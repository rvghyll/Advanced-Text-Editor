import tkinter as tk
from tkinter import filedialog, messagebox, ttk, colorchooser, simpledialog
from tkinter.font import Font
from PIL import Image, ImageTk, ImageDraw
import os
import tempfile
import json
import time
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

try:
    from openpyxl import Workbook
    from openpyxl.styles import Font as ExcelFont, Alignment
    excel_available = True
except ImportError:
    excel_available = False

spell_checker = None
personal_dictionary = set()

try:
    import enchant
    spell_checker = enchant.Dict("en_US")
    # Load personal dictionary from file
    personal_dict_file = os.path.join(tempfile.gettempdir(), "notepad_personal_dict.txt")
    if os.path.exists(personal_dict_file):
        with open(personal_dict_file, 'r') as f:
            personal_dictionary = set(line.strip() for line in f if line.strip())
except ImportError:
    pass

def save_personal_dictionary():
    if spell_checker:
        personal_dict_file = os.path.join(tempfile.gettempdir(), "notepad_personal_dict.txt")
        with open(personal_dict_file, 'w') as f:
            for word in sorted(personal_dictionary):
                f.write(word + '\n')

# --------------------------
# Tab Data Class
# --------------------------
class Tab:
    def __init__(self, name="Untitled", is_drawing=False):
        self.name = name
        self.is_drawing = is_drawing
        self.content = None # This will hold either Text widget or DrawingCanvas object
        self.file_path = None
        self.is_modified = False
        self.images = [] # For text tabs with inserted images
        self.image_paths = [] # Paths of inserted images
        self.last_edit_time = ""
        self.last_save_time = ""
        self.theme = "light" # Default theme for text tabs
        self.auto_save_path = None # For drawing tabs, path where drawing is saved
        self.drawing_text = "" # Associated text for drawing tabs
        self.linked_text_tab_name = "" # Name of linked text tab
        self.linked_drawing_tab_name = "" # For text tabs: name of linked drawing tab

# --------------------------
# Drawing Canvas Class
# --------------------------
class DrawingCanvas:
    def __init__(self, parent, app, dark_mode=False):
        self.app = app
        self.parent = parent
        self.dark_mode = dark_mode
        self.drawing = False
        self.brush_size = 3
        self.brush_color = "#FFFFFF" if dark_mode else "#000000"
        self.is_eraser = False
        self.last_x = 0
        self.last_y = 0
        self.start_x = 0
        self.start_y = 0
        self.canvas_theme = "light" # Theme for the canvas background

        self.frame = tk.Frame(parent)

        self.toolbar_frame = tk.Frame(self.frame)
        self.toolbar_frame.pack(side=tk.TOP, fill=tk.X, padx=5, pady=5)

        tk.Label(self.toolbar_frame, text="Brush Size:", bg="#f3f3f3" if not self.dark_mode else "#252526", fg="#000000" if not self.dark_mode else "#ffffff").pack(side=tk.LEFT, padx=5)
        self.size_slider = tk.Scale(self.toolbar_frame, from_=1, to=50, orient=tk.HORIZONTAL, length=150, bg="#f3f3f3" if not self.dark_mode else "#252526", troughcolor="#cccccc" if not self.dark_mode else "#555555")
        self.size_slider.set(self.brush_size)
        self.size_slider.pack(side=tk.LEFT, padx=5)
        self.size_slider.bind("<Motion>", lambda e: self.on_brush_size_changed())
        self.size_slider.bind("<Button-1>", lambda e: self.on_brush_size_changed())

        self.size_value_label = tk.Label(self.toolbar_frame, text=f"{self.brush_size}px", width=4, bg="#f3f3f3" if not self.dark_mode else "#252526", fg="#000000" if not self.dark_mode else "#ffffff")
        self.size_value_label.pack(side=tk.LEFT, padx=5)

        tk.Label(self.toolbar_frame, text="  ").pack(side=tk.LEFT) # Spacer

        self.color_btn = tk.Button(self.toolbar_frame, text="Color", command=self.pick_color, width=8, bg="#f3f3f3" if not self.dark_mode else "#252526", fg="#000000" if not self.dark_mode else "#ffffff")
        self.color_btn.pack(side=tk.LEFT, padx=5)

        self.eraser_btn = tk.Button(self.toolbar_frame, text="Eraser", command=self.toggle_eraser, width=8, bg="#f3f3f3" if not self.dark_mode else "#252526", fg="#000000" if not self.dark_mode else "#ffffff")
        self.eraser_btn.pack(side=tk.LEFT, padx=5)

        tk.Label(self.toolbar_frame, text="  ").pack(side=tk.LEFT) # Spacer

        tk.Label(self.toolbar_frame, text="Mode:", bg="#f3f3f3" if not self.dark_mode else "#252526", fg="#000000" if not self.dark_mode else "#ffffff").pack(side=tk.LEFT, padx=5)
        self.mode_var = tk.StringVar(value="pen")
        modes = [("Pen", "pen"), ("Line", "line"), ("Rect", "rect"), ("Oval", "oval")]
        for text, val in modes:
            b = tk.Radiobutton(self.toolbar_frame, text=text, variable=self.mode_var, value=val, indicatoron=False, width=6,
                               bg="#f3f3f3" if not self.dark_mode else "#252526", fg="#000000" if not self.dark_mode else "#ffffff",
                               activebackground="#cccccc" if not self.dark_mode else "#666666",
                               selectcolor="#a0a0a0" if not self.dark_mode else "#777777")
            b.pack(side=tk.LEFT, padx=2)

        clear_btn = tk.Button(self.toolbar_frame, text="Clear All", command=self.clear_canvas, width=10, bg="#f3f3f3" if not self.dark_mode else "#252526", fg="#000000" if not self.dark_mode else "#ffffff")
        clear_btn.pack(side=tk.LEFT, padx=5)

        self.canvas = tk.Canvas(self.frame, bg="white", cursor="cross")
        self.canvas.pack(fill=tk.BOTH, expand=True)

        self.canvas.bind("<Button-1>", self.on_mouse_press)
        self.canvas.bind("<B1-Motion>", self.on_mouse_move)
        self.canvas.bind("<ButtonRelease-1>", self.on_mouse_release)

        self.canvas_width = 1200
        self.canvas_height = 900
        img_bg_color = (30, 30, 30) if self.dark_mode else (255, 255, 255)
        self.image = Image.new("RGB", (self.canvas_width, self.canvas_height), color=img_bg_color)

        self.set_dark_mode(self.dark_mode)

    def on_brush_size_changed(self):
        self.brush_size = self.size_slider.get()
        self.size_value_label.config(text=f"{self.brush_size}px")

    def pick_color(self):
        color = colorchooser.askcolor(color=self.brush_color)[1]
        if color:
            self.brush_color = color

    def toggle_eraser(self):
        self.is_eraser = not self.is_eraser
        self.eraser_btn.config(bg="#ffeb3b" if self.is_eraser else ("#f3f3f3" if not self.dark_mode else "#252526"))

    def clear_canvas(self):
        bg_color = (30, 30, 30) if self.dark_mode else (255, 255, 255)
        self.image = Image.new("RGB", (self.canvas_width, self.canvas_height), color=bg_color)
        self.canvas.delete("all")
        self.app.log_action("Cleared drawing canvas")

    def set_dark_mode(self, dark_mode):
        self.dark_mode = dark_mode
        bg_color = "#1e1e1e" if dark_mode else "#ffffff"
        self.canvas.config(bg=bg_color)
        img_bg_color = (30, 30, 30) if dark_mode else (255, 255, 255)
        self.image = Image.new("RGB", (self.canvas_width, self.canvas_height), color=img_bg_color)
        self.canvas.delete("all")
        
        self.brush_color = "#FFFFFF" if dark_mode else "#000000"
        
        # Update toolbar colors
        self.toolbar_frame.config(bg="#252526" if dark_mode else "#f3f3f3")
        for widget in self.toolbar_frame.winfo_children():
            if isinstance(widget, tk.Label):
                widget.config(bg="#252526" if dark_mode else "#f3f3f3", fg="#ffffff" if dark_mode else "#000000")
            elif isinstance(widget, tk.Button):
                widget.config(bg="#252526" if dark_mode else "#f3f3f3", fg="#ffffff" if dark_mode else "#000000")
            elif isinstance(widget, tk.Scale):
                widget.config(bg="#252526" if dark_mode else "#f3f3f3", troughcolor="#555555" if dark_mode else "#cccccc")
            elif isinstance(widget, tk.Radiobutton):
                widget.config(bg="#252526" if dark_mode else "#f3f3f3", fg="#ffffff" if dark_mode else "#000000",
                              activebackground="#666666" if dark_mode else "#cccccc",
                              selectcolor="#777777" if dark_mode else "#a0a0a0")

    def set_theme(self, theme):
        self.canvas_theme = theme
        theme_colors = {
            "light": "#ffffff",
            "dark": "#1e1e1e",
            "sepia": "#f4eae0",
            "blue": "#e3f2fd",
            "green": "#e8f5e9"
        }
        bg_color = theme_colors.get(theme, "#ffffff")
        self.canvas.config(bg=bg_color)
        
        rgb_colors = {
            "light": (255, 255, 255),
            "dark": (30, 30, 30),
            "sepia": (244, 234, 224),
            "blue": (227, 242, 253),
            "green": (232, 245, 233)
        }
        img_bg_color = rgb_colors.get(theme, (255, 255, 255))
        self.image = Image.new("RGB", (self.canvas_width, self.canvas_height), color=img_bg_color)
        self.canvas.delete("all")
        self.app.log_action(f"Changed drawing theme to {theme}")

    def on_mouse_press(self, event):
        self.drawing = True
        self.start_x = event.x
        self.start_y = event.y
        self.last_x = event.x
        self.last_y = event.y
        self.canvas.delete("preview")
        self.app.log_action("Drawing: press")

    def on_mouse_move(self, event):
        if not self.drawing:
            return
        x, y = event.x, event.y
        mode = self.mode_var.get()
        draw = ImageDraw.Draw(self.image)

        if mode == "pen":
            # Calculate color based on eraser and dark mode
            if self.is_eraser:
                color_rgb = (30, 30, 30) if self.dark_mode else (255, 255, 255)
            else:
                color_rgb = self.hex_to_rgb(self.brush_color)
            
            draw.line([(self.last_x, self.last_y), (x, y)], fill=color_rgb, width=self.brush_size)
            
            # Canvas drawing color needs to consider eraser for black/white background
            canvas_draw_color = ("#1e1e1e" if self.dark_mode else "#ffffff") if self.is_eraser else self.brush_color
            self.canvas.create_line(self.last_x, self.last_y, x, y, width=self.brush_size,
                                    fill=canvas_draw_color,
                                    capstyle=tk.ROUND, smooth=True)
            self.last_x = x
            self.last_y = y
        else:
            self.canvas.delete("preview")
            if mode == "line":
                self.canvas.create_line(self.start_x, self.start_y, x, y, width=self.brush_size, fill=self.brush_color, tags="preview", capstyle=tk.ROUND)
            elif mode == "rect":
                self.canvas.create_rectangle(self.start_x, self.start_y, x, y, width=self.brush_size, outline=self.brush_color, tags="preview")
            elif mode == "oval":
                self.canvas.create_oval(self.start_x, self.start_y, x, y, width=self.brush_size, outline=self.brush_color, tags="preview")

    def on_mouse_release(self, event):
        if not self.drawing:
            return
        self.drawing = False
        x, y = event.x, event.y
        mode = self.mode_var.get()
        draw = ImageDraw.Draw(self.image)
        
        # Calculate color for image drawing based on eraser and dark mode
        if self.is_eraser:
            color_rgb = (30, 30, 30) if self.dark_mode else (255, 255, 255)
        else:
            color_rgb = self.hex_to_rgb(self.brush_color)

        if mode == "pen":
            # Pen drawing is handled in on_mouse_move, no need to do anything here except set drawing to False
            pass
        else:
            if mode == "line":
                draw.line([(self.start_x, self.start_y), (x, y)], fill=color_rgb, width=self.brush_size)
                self.canvas.create_line(self.start_x, self.start_y, x, y, width=self.brush_size, fill=self.brush_color, capstyle=tk.ROUND)
            elif mode == "rect":
                draw.rectangle([(self.start_x, self.start_y), (x, y)], outline=color_rgb, width=self.brush_size)
                self.canvas.create_rectangle(self.start_x, self.start_y, x, y, width=self.brush_size, outline=self.brush_color)
            elif mode == "oval":
                draw.ellipse([(self.start_x, self.start_y), (x, y)], outline=color_rgb, width=self.brush_size)
                self.canvas.create_oval(self.start_x, self.start_y, x, y, width=self.brush_size, outline=self.brush_color)

        self.canvas.delete("preview")
        self.app.log_action(f"Drawing: {mode} placed")

    def hex_to_rgb(self, hex_color):
        hex_color = hex_color.lstrip('#')
        return tuple(int(hex_color[i:i+2], 16) for i in (0, 2, 4))

    def save_image(self, file_path):
        self.image.save(file_path)

# --------------------------
# Notepad Application Class
# --------------------------
class NotepadApp:
    def __init__(self, root):
        self.root = root
        self.root.title("Notepad with Drawing")
        self.root.geometry("1200x800")

        self.tabs = []
        self.current_tab_index = 0
        self.dark_mode = False
        self.bold = False
        self.italic = False
        self.underline = False
        self.used_text_tab_numbers = set()
        self.used_drawing_tab_numbers = set()
        self.highlight_tags = {} # To manage custom highlight tags
        self.spell_check_enabled = True if spell_checker else False
        self.last_spell_check_word = None
        self.activity_log = []
        self.session_file = os.path.join(tempfile.gettempdir(), "notepad_session.json")

        self.main_frame = tk.Frame(root)
        self.main_frame.pack(fill=tk.BOTH, expand=True)

        self.create_menu_bar()
        self.create_toolbar()

        self.notebook = ttk.Notebook(self.main_frame)
        self.notebook.pack(fill=tk.BOTH, expand=True, side=tk.TOP)
        self.notebook.bind("<Button-3>", self.show_tab_context_menu)
        self.notebook.bind("<<NotebookTabChanged>>", self.on_tab_change)

        self.enable_tab_reordering()

        self.new_text_tab()
        self.load_session()

        self.root.bind("<Control-n>", lambda e: self.new_text_tab())
        self.root.bind("<Control-o>", lambda e: self.open_file())
        self.root.bind("<Control-s>", lambda e: self.save_file())
        self.root.bind("<Control-Shift-S>", lambda e: self.save_file_as())
        self.root.bind("<Control-p>", lambda e: self.print_file())
        self.root.bind("<Control-q>", lambda e: self.root.quit())
        self.root.bind("<Control-f>", lambda e: self.search_text())
        self.root.bind("<Control-b>", lambda e: self.toggle_bold())
        self.root.bind("<Control-i>", lambda e: self.toggle_italic())
        self.root.bind("<Control-u>", lambda e: self.toggle_underline())

        self.root.protocol("WM_DELETE_WINDOW", self.on_close)
        self.root.after(30000, self.autosave) # Autosave every 30 seconds

    def create_menu_bar(self):
        menubar = tk.Menu(self.root)
        self.root.config(menu=menubar)

        file_menu = tk.Menu(menubar, tearoff=0)
        menubar.add_cascade(label="File", menu=file_menu)
        file_menu.add_command(label="New Text Tab", command=self.new_text_tab, accelerator="Ctrl+N")
        file_menu.add_command(label="New Drawing Tab", command=self.new_drawing_tab)
        file_menu.add_separator()
        file_menu.add_command(label="Open", command=self.open_file, accelerator="Ctrl+O")
        file_menu.add_command(label="Save", command=self.save_file, accelerator="Ctrl+S")
        file_menu.add_command(label="Save As...", command=self.save_file_as, accelerator="Ctrl+Shift+S")
        file_menu.add_separator()
        file_menu.add_command(label="Print", command=self.print_file, accelerator="Ctrl+P")
        file_menu.add_separator()
        file_menu.add_command(label="Export PDF (with images)", command=self.export_pdf_with_images)
        file_menu.add_command(label="Export Word (with images)", command=self.export_word_with_images)
        file_menu.add_command(label="Export to Excel", command=self.export_excel)
        file_menu.add_separator()
        file_menu.add_command(label="Export Drawing PDF", command=self.export_drawing_pdf)
        file_menu.add_command(label="Export Drawing Word", command=self.export_drawing_word)
        file_menu.add_separator()
        file_menu.add_command(label="Insert Image", command=self.insert_image)
        file_menu.add_command(label="Export Images Separately", command=self.export_images)
        file_menu.add_separator()
        file_menu.add_command(label="Exit", command=self.on_close, accelerator="Ctrl+Q")

        edit_menu = tk.Menu(menubar, tearoff=0)
        menubar.add_cascade(label="Edit", menu=edit_menu)
        edit_menu.add_command(label="Undo", command=self.undo, accelerator="Ctrl+Z")
        edit_menu.add_command(label="Redo", command=self.redo, accelerator="Ctrl+Y")
        edit_menu.add_separator()
        edit_menu.add_command(label="Cut", command=self.cut, accelerator="Ctrl+X")
        edit_menu.add_command(label="Copy", command=self.copy, accelerator="Ctrl+C")
        edit_menu.add_command(label="Paste", command=self.paste, accelerator="Ctrl+V")
        edit_menu.add_separator()
        edit_menu.add_command(label="Delete", command=self.delete_selection, accelerator="Delete")
        edit_menu.add_separator()
        edit_menu.add_command(label="Select All", command=self.select_all, accelerator="Ctrl+A")
        edit_menu.add_separator()
        edit_menu.add_command(label="Search", command=self.search_text, accelerator="Ctrl+F")
        edit_menu.add_command(label="Export Image with Current Text", command=self.export_image_with_current_text)
        if spell_checker:
            edit_menu.add_command(label="Spell Check", command=self.spell_check)

        view_menu = tk.Menu(menubar, tearoff=0)
        menubar.add_cascade(label="View", menu=view_menu)
        view_menu.add_command(label="Dark Mode", command=self.toggle_dark_mode)
        view_menu.add_command(label="Choose Theme", command=self.pick_theme)

        tools_menu = tk.Menu(menubar, tearoff=0)
        menubar.add_cascade(label="Tools", menu=tools_menu)
        tools_menu.add_command(label="Activity Log", command=self.open_activity_log)
        if spell_checker:
            tools_menu.add_command(label="Manage Personal Dictionary", command=self.manage_personal_dictionary)

        help_menu = tk.Menu(menubar, tearoff=0)
        menubar.add_cascade(label="Help", menu=help_menu)
        help_menu.add_command(label="About", command=self.show_about)
        
        self.status_label = tk.Label(menubar, text="Words: 0 | Characters: 0 | Lines: 1 | Column: 1", font=("Arial", 9), bg="#f0f0f0", padx=10)

    def create_toolbar(self):
        toolbar = tk.Frame(self.root, bg="#f3f3f3", relief=tk.RAISED, height=40)
        toolbar.pack(side=tk.TOP, fill=tk.X)
        toolbar.pack_propagate(False) # Prevent frame from shrinking to fit content

        tk.Label(toolbar, text="Font:", bg="#f3f3f3").pack(side=tk.LEFT, padx=5)
        self.font_var = tk.StringVar(value="Arial")
        # Common fonts list
        font_families = sorted(tk.font.families())
        self.font_combo = ttk.Combobox(toolbar, textvariable=self.font_var, values=font_families, state="readonly", width=18)
        self.font_combo.pack(side=tk.LEFT, padx=5)
        self.font_combo.bind("<<ComboboxSelected>>", lambda e: self.change_font())

        tk.Label(toolbar, text="Size:", bg="#f3f3f3").pack(side=tk.LEFT, padx=5)
        self.size_var = tk.StringVar(value="10")
        # Common font sizes, can be extended
        font_sizes = [str(i) for i in range(8, 73, 2)] # From 8 to 72, step 2
        self.size_combo = ttk.Combobox(toolbar, textvariable=self.size_var, values=font_sizes, state="readonly", width=5)
        self.size_combo.pack(side=tk.LEFT, padx=5)
        self.size_combo.bind("<<ComboboxSelected>>", lambda e: self.change_font_size())

        self.bold_btn = tk.Button(toolbar, text="B", command=self.toggle_bold, bg="#f3f3f3", font=("Arial", 10, "bold"), width=3)
        self.bold_btn.pack(side=tk.LEFT, padx=2)

        self.italic_btn = tk.Button(toolbar, text="I", command=self.toggle_italic, bg="#f3f3f3", font=("Arial", 10, "italic"), width=3)
        self.italic_btn.pack(side=tk.LEFT, padx=2)

        self.underline_btn = tk.Button(toolbar, text="U", command=self.toggle_underline, bg="#f3f3f3", font=("Arial", 10, "underline"), width=3)
        self.underline_btn.pack(side=tk.LEFT, padx=2)
        
        self.highlight_btn = tk.Button(toolbar, text="H", command=self.show_highlight_palette, bg="#f3f3f3", font=("Arial", 10, "bold"), width=3)
        self.highlight_btn.pack(side=tk.LEFT, padx=2)
        
        if self.spell_check_enabled:
            self.spell_check_btn = tk.Button(toolbar, text="Spell Check", command=self.spell_check, bg="#f3f3f3", width=10)
            self.spell_check_btn.pack(side=tk.LEFT, padx=5)

        # Add hover effects for buttons
        for btn in (self.bold_btn, self.italic_btn, self.underline_btn, self.highlight_btn, self.spell_check_btn if self.spell_check_enabled else None):
            if btn:
                btn.bind("<Enter>", lambda e, b=btn: b.config(relief=tk.RAISED))
                btn.bind("<Leave>", lambda e, b=btn: b.config(relief=tk.FLAT))
            
        # Status bar at the right end of the toolbar
        self.toolbar_status = tk.Label(toolbar, text="Words: 0 | Characters: 0 | Lines: 1 | Column: 1", bg="#f3f3f3", font=("Arial", 9))
        self.toolbar_status.pack(side=tk.RIGHT, padx=10)

    def get_next_text_tab_number(self):
        """Finds the next available integer for a text tab number."""
        if not self.used_text_tab_numbers:
            return 1
        num = 1
        while num in self.used_text_tab_numbers:
            num += 1
        return num

    def get_next_drawing_tab_number(self):
        """Finds the next available integer for a drawing tab number."""
        if not self.used_drawing_tab_numbers:
            return 1
        num = 1
        while num in self.used_drawing_tab_numbers:
            num += 1
        return num

    def new_text_tab(self):
        tab_number = self.get_next_text_tab_number()
        self.used_text_tab_numbers.add(tab_number)
        tab = Tab(f"Untitled {tab_number}", False)
        tab.tab_number = tab_number # Assign a unique number to the tab
        self.tabs.append(tab)

        frame = tk.Frame(self.notebook)
        text_widget = tk.Text(frame, wrap=tk.WORD, undo=True, font=("Arial", 10))
        text_widget.pack(fill=tk.BOTH, expand=True)
        
        # Bindings for text widget
        text_widget.bind("<KeyRelease>", lambda e: (self.update_status_bar(e), self.highlight_syntax(text_widget), self.check_spelling_realtime(text_widget)))
        text_widget.bind("<<Modified>>", lambda e: text_widget.edit_modified(False)) # Prevent excessive "Modified" events
        text_widget.bind("<Button-3>", lambda e: self.show_spell_context_menu(e, text_widget)) # Right-click for spell check context menu
        text_widget.tag_configure("spell_error", underline=True, underlinefg="red") # Style for misspelled words

        tab.content = text_widget
        self.notebook.add(frame, text=tab.name)
        self.notebook.select(len(self.tabs)-1) # Select the newly created tab
        self.current_tab_index = len(self.tabs)-1 # Update current tab index

        self.apply_font_style(tab) # Apply default font settings
        self.update_style_buttons() # Ensure style buttons reflect default state
        self.log_action(f"Created new text tab: {tab.name}")

    def new_drawing_tab(self):
        tab_number = self.get_next_drawing_tab_number()
        self.used_drawing_tab_numbers.add(tab_number)
        tab = Tab(f"Drawing {tab_number}", True)
        tab.tab_number = tab_number # Assign a unique number to the tab
        self.tabs.append(tab)

        frame = tk.Frame(self.notebook)
        drawing_canvas = DrawingCanvas(frame, self, self.dark_mode)
        drawing_canvas.frame.pack(fill=tk.BOTH, expand=True)

        tab.content = drawing_canvas
        self.notebook.add(frame, text=tab.name)
        self.notebook.select(len(self.tabs)-1) # Select the newly created tab
        self.current_tab_index = len(self.tabs)-1 # Update current tab index
        self.log_action(f"Created new drawing tab: {tab.name}")

    def on_tab_change(self, event):
        """Handles tab changes, updating current tab index and status bar."""
        try:
            self.current_tab_index = self.notebook.index(self.notebook.select())
            self.update_status_bar() # Update status bar based on the new tab
            self.update_style_buttons() # Update style buttons (bold, italic, etc.)
        except tk.TclError:
            # This can happen if all tabs are closed, ignore
            pass

    def show_tab_context_menu(self, event):
        """Shows a context menu when right-clicking on a tab."""
        try:
            tab_index = self.notebook.index(f"@{event.x},{event.y}")
            menu = tk.Menu(self.root, tearoff=0)
            menu.add_command(label="Rename Tab", command=lambda: self.rename_tab(tab_index))
            menu.add_command(label="Close Tab", command=lambda: self.close_tab(tab_index))
            menu.post(event.x_root, event.y_root)
        except tk.TclError:
            # Clicked outside of a tab, do nothing
            pass

    def rename_tab(self, index):
        """Renames the tab at the given index."""
        if index < len(self.tabs):
            new_name = simpledialog.askstring("Rename Tab", "Enter new tab name:", initialvalue=self.tabs[index].name)
            if new_name:
                self.tabs[index].name = new_name
                self.notebook.tab(index, text=new_name)
                self.log_action(f"Renamed tab to: {new_name}")

    def close_tab(self, index):
        """Closes the tab at the given index."""
        if index < len(self.tabs):
            tab = self.tabs[index]
            # Remove the tab number from the used sets
            if hasattr(tab, 'tab_number'):
                if tab.is_drawing:
                    self.used_drawing_tab_numbers.discard(tab.tab_number)
                else:
                    self.used_text_tab_numbers.discard(tab.tab_number)
                
            self.notebook.forget(index) # Remove tab from notebook widget
            self.tabs.pop(index)       # Remove tab from our internal list
            self.log_action(f"Closed tab: {tab.name}")
            
            # If no tabs left, create a new one
            if not self.tabs:
                self.new_text_tab()
            else:
                # Update current_tab_index to the previous tab or the first if the removed was the first
                self.current_tab_index = max(0, index - 1)
                self.on_tab_change(None) # Update UI based on new current tab

    def open_file(self):
        """Opens a text file and loads its content into a new tab."""
        file_path = filedialog.askopenfilename(filetypes=[("Text Files", "*.txt"), ("Markdown", "*.md"), ("All Files", "*.*")])
        if not file_path:
            return
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                content = f.read()

            tab_number = self.get_next_text_tab_number()
            self.used_text_tab_numbers.add(tab_number)
            tab = Tab(os.path.basename(file_path), False)
            tab.file_path = file_path
            tab.tab_number = tab_number
            self.tabs.append(tab)

            frame = tk.Frame(self.notebook)
            text_widget = tk.Text(frame, wrap=tk.WORD, undo=True)
            text_widget.pack(fill=tk.BOTH, expand=True)
            text_widget.bind("<KeyRelease>", lambda e: (self.update_status_bar(e), self.highlight_syntax(text_widget)))
            text_widget.insert("1.0", content)

            tab.content = text_widget
            self.notebook.add(frame, text=tab.name)
            self.notebook.select(len(self.tabs)-1)
            self.current_tab_index = len(self.tabs)-1

            self.apply_font_style(tab) # Apply default font style
            self.log_action(f"Opened file: {file_path}")
        except Exception as e:
            messagebox.showerror("Open File", f"Failed to open: {e}")

    def save_file(self):
        """Saves the current tab's content."""
        if not self.tabs: return
        tab = self.tabs[self.current_tab_index]

        if tab.is_drawing:
            # For drawing tabs, try to save to an existing path if available
            if hasattr(tab, 'auto_save_path') and tab.auto_save_path:
                try:
                    tab.content.save_image(tab.auto_save_path)
                    tab.last_save_time = time.strftime("%Y-%m-%d %H:%M:%S")
                    self.log_action(f"Saved drawing to {tab.auto_save_path}")
                    messagebox.showinfo("Save", f"Drawing saved: {tab.auto_save_path}")
                    return # Exit after saving to auto_save_path
                except Exception as e:
                    messagebox.showerror("Save", f"Failed to save drawing: {e}")
            
            # If no auto_save_path, prompt for save as
            file_path = filedialog.asksaveasfilename(defaultextension=".png", filetypes=[("PNG File","*.png"), ("JPEG File","*.jpg")])
            if not file_path:
                return
            try:
                tab.content.save_image(file_path)
                tab.auto_save_path = file_path # Store path for future saves
                tab.last_save_time = time.strftime("%Y-%m-%d %H:%M:%S")
                self.log_action(f"Saved drawing to {file_path}")
                messagebox.showinfo("Save", f"Drawing saved: {file_path}")
            except Exception as e:
                messagebox.showerror("Save", f"Failed to save drawing: {e}")
            return

        # For text tabs
        if tab.file_path: # If file has been saved before
            try:
                with open(tab.file_path, "w", encoding="utf-8") as f:
                    f.write(tab.content.get("1.0", tk.END))
                tab.last_save_time = time.strftime("%Y-%m-%d %H:%M:%S")
                tab.is_modified = False
                self.log_action(f"Saved file: {tab.file_path}")
                self.update_status_bar() # Update status bar to reflect no modification
            except Exception as e:
                messagebox.showerror("Save", f"Failed to save: {e}")
        else: # If file has never been saved
            self.save_file_as()

    def save_file_as(self):
        """Saves the current tab's content to a new file path."""
        if not self.tabs: return
        tab = self.tabs[self.current_tab_index]

        if tab.is_drawing:
            file_path = filedialog.asksaveasfilename(defaultextension=".png", filetypes=[("PNG File","*.png"), ("JPEG File","*.jpg")])
            if not file_path:
                return
            try:
                tab.content.save_image(file_path)
                tab.auto_save_path = file_path # Update auto_save_path
                tab.last_save_time = time.strftime("%Y-%m-%d %H:%M:%S")
                self.log_action(f"Saved drawing to {file_path}")
                messagebox.showinfo("Save As", f"Drawing saved: {file_path}")
            except Exception as e:
                messagebox.showerror("Save As", f"Failed to save: {e}")
            return

        # For text tabs
        file_path = filedialog.asksaveasfilename(defaultextension=".txt", filetypes=[("Text Files","*.txt"),("Markdown","*.md")])
        if not file_path:
            return
        try:
            with open(file_path, "w", encoding="utf-8") as f:
                f.write(tab.content.get("1.0", tk.END))
            tab.file_path = file_path
            tab.name = os.path.basename(file_path)
            self.notebook.tab(self.current_tab_index, text=tab.name) # Update tab title
            tab.last_save_time = time.strftime("%Y-%m-%d %H:%M:%S")
            tab.is_modified = False
            self.log_action(f"Saved file as: {file_path}")
        except Exception as e:
            messagebox.showerror("Save As", f"Failed to save: {e}")

    def insert_image(self):
        """Inserts an image into the current text tab."""
        if not self.tabs or self.current_tab_index >= len(self.tabs):
            return
        tab = self.tabs[self.current_tab_index]
        if tab.is_drawing:
            messagebox.showinfo("Insert Image", "Use drawing tools for drawing tabs")
            return

        file_path = filedialog.askopenfilename(filetypes=[("Image Files", "*.png *.jpg *.jpeg *.gif"), ("All Files", "*.*")])
        if not file_path:
            return
        try:
            img = Image.open(file_path)
            img.thumbnail((400, 300)) # Resize image if too large
            photo = ImageTk.PhotoImage(img)
            
            # Insert image marker in text widget
            tab.content.insert(tk.INSERT, "\n[IMAGE]\n")
            
            # Store the image path and PhotoImage reference
            tab.image_paths.append(file_path)
            tab.images.append(photo) # Keep reference to prevent garbage collection
            
            self.log_action(f"Inserted image: {file_path}")
            messagebox.showinfo("Image Inserted", 
                              "Image marker inserted. The image will be included when exporting to PDF or Word.")
        except Exception as e:
            messagebox.showerror("Insert Image", f"Failed: {e}")

    def export_images(self):
        """Exports all inserted images from a text tab to a selected directory."""
        if not self.tabs or self.current_tab_index >= len(self.tabs):
            return
        tab = self.tabs[self.current_tab_index]
        if tab.is_drawing:
            messagebox.showinfo("Export Images", "This function is for text tabs with inserted images.")
            return
        
        if not tab.image_paths:
            messagebox.showinfo("Export Images", "No images have been inserted in this tab.")
            return
        
        export_dir = filedialog.askdirectory(title="Select folder to export images")
        if not export_dir:
            return
        
        try:
            exported_count = 0
            for i, img_path in enumerate(tab.image_paths):
                if os.path.exists(img_path):
                    # Try to copy the image file
                    try:
                        img = Image.open(img_path)
                        # Construct a safe filename
                        base_name = os.path.basename(img_path)
                        # Ensure filename is unique and has an extension
                        name, ext = os.path.splitext(base_name)
                        if not name: name = "image"
                        if not ext: ext = ".png" # Default to png if no extension
                        
                        filename = f"exported_{i+1}_{name}{ext}"
                        save_path = os.path.join(export_dir, filename)
                        
                        # Check for duplicates and append number if necessary
                        counter = 1
                        while os.path.exists(save_path):
                            filename = f"exported_{i+1}_{name}_{counter}{ext}"
                            save_path = os.path.join(export_dir, filename)
                            counter += 1
                        
                        img.save(save_path)
                        exported_count += 1
                    except Exception as img_e:
                        print(f"Warning: Could not export image {img_path}: {img_e}")
            
            if exported_count > 0:
                messagebox.showinfo("Export Images", f"Exported {exported_count} image(s) to:\n{export_dir}")
                self.log_action(f"Exported {exported_count} images")
            else:
                messagebox.showinfo("Export Images", "No images were successfully exported.")
        except Exception as e:
            messagebox.showerror("Export Images", f"An error occurred: {e}")

    def export_pdf_with_images(self):
        """Exports the current text tab content with inserted images to a PDF file."""
        if not self.tabs: return
        tab = self.tabs[self.current_tab_index]
        if tab.is_drawing:
            messagebox.showinfo("Export PDF", "Use 'Export Drawing PDF' for drawing tabs.")
            return
        
        file_path = filedialog.asksaveasfilename(defaultextension=".pdf", filetypes=[("PDF Files","*.pdf")])
        if not file_path:
            return
        
        try:
            c = canvas.Canvas(file_path, pagesize=letter)
            text_content = tab.content.get("1.0", tk.END)
            
            # Split content by lines
            lines = text_content.split("\n")
            
            y_position = 750 # Starting y position
            line_height = 15 # Spacing between lines
            max_chars_per_line = 90 # Approximate max characters before wrapping
            
            image_index = 0
            page_number = 1
            
            # Add title
            c.setFont("Helvetica-Bold", 16)
            c.drawString(50, 780, f"Document: {tab.name}")
            y_position = 750
            
            for line in lines:
                # Check for image marker
                if line.strip() == "[IMAGE]" and image_index < len(tab.image_paths):
                    # Try to add the image
                    try:
                        img_path = tab.image_paths[image_index]
                        if os.path.exists(img_path):
                            img = Image.open(img_path)
                            
                            # Check if we need a new page for the image
                            if y_position < 200: # Need space for image
                                c.showPage()
                                y_position = 750
                                page_number += 1
                            
                            # Add image caption
                            c.setFont("Helvetica", 10)
                            c.drawString(50, y_position, f"Image {image_index + 1}:")
                            y_position -= 15
                            
                            # Calculate image dimensions
                            img_width, img_height = img.size
                            max_width = 500
                            max_height = 300
                            
                            # Maintain aspect ratio
                            if img_width > max_width or img_height > max_height:
                                ratio = min(max_width/img_width, max_height/img_height)
                                img_width = int(img_width * ratio)
                                img_height = int(img_height * ratio)
                            
                            # Center image on page
                            x_offset = (letter[0] - img_width) / 2
                            
                            # Save image to temporary file for PDF
                            temp_img_path = os.path.join(tempfile.gettempdir(), f"temp_pdf_img_{image_index}.png")
                            img.save(temp_img_path)
                            
                            # Draw image
                            c.drawImage(temp_img_path, x_offset, y_position - img_height, 
                                       width=img_width, height=img_height)
                            
                            # Update position
                            y_position -= (img_height + 20)
                            
                            # Clean up temporary file
                            try:
                                os.remove(temp_img_path)
                            except:
                                pass
                            
                            image_index += 1
                        else:
                            # Image not found, skip
                            c.setFont("Helvetica", 10)
                            c.drawString(50, y_position, f"[Image {image_index + 1} not found]")
                            y_position -= line_height
                            image_index += 1
                    except Exception as img_e:
                        print(f"Warning: Could not add image to PDF: {img_e}")
                        c.setFont("Helvetica", 10)
                        c.drawString(50, y_position, f"[Error loading image {image_index + 1}]")
                        y_position -= line_height
                        image_index += 1
                    continue
                
                # Handle regular text lines
                if line.strip() == "":
                    # Empty line
                    if y_position < 50:
                        c.showPage()
                        y_position = 750
                        page_number += 1
                    else:
                        y_position -= line_height
                    continue
                
                # Basic manual wrapping for long lines
                while len(line) > max_chars_per_line:
                    if y_position < 50:
                        c.showPage()
                        y_position = 750
                        page_number += 1
                    
                    c.setFont("Helvetica", 12)
                    c.drawString(50, y_position, line[:max_chars_per_line])
                    line = line[max_chars_per_line:]
                    y_position -= line_height
                
                if y_position < 50:
                    c.showPage()
                    y_position = 750
                    page_number += 1
                
                c.setFont("Helvetica", 12)
                c.drawString(50, y_position, line)
                y_position -= line_height
            
            # Add page number on last page
            c.setFont("Helvetica", 8)
            c.drawString(550, 30, f"Page {page_number}")
            
            c.save()
            messagebox.showinfo("Export Successful", 
                              f"Document with images exported to:\n{file_path}\n\n"
                              f"Total images included: {image_index}")
            self.log_action(f"Exported PDF with images: {file_path}")
            
        except Exception as e:
            messagebox.showerror("Export PDF", f"Failed to export: {e}")

    def export_word_with_images(self):
        """Exports the current text tab content with inserted images to a Word (.docx) file."""
        if not self.tabs: return
        tab = self.tabs[self.current_tab_index]
        if tab.is_drawing:
            messagebox.showinfo("Export Word", "Use 'Export Drawing Word' for drawing tabs.")
            return
        
        file_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word Files","*.docx")])
        if not file_path:
            return
        
        try:
            doc = Document()
            
            # Add title
            title = doc.add_heading(f"Document: {tab.name}", 0)
            title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            
            # Get text content
            text_content = tab.content.get("1.0", tk.END)
            
            # Split content by lines
            lines = text_content.split("\n")
            
            image_index = 0
            
            for line in lines:
                # Check for image marker
                if line.strip() == "[IMAGE]" and image_index < len(tab.image_paths):
                    # Try to add the image
                    try:
                        img_path = tab.image_paths[image_index]
                        if os.path.exists(img_path):
                            # Add image caption
                            p = doc.add_paragraph()
                            run = p.add_run(f"Image {image_index + 1}:")
                            run.bold = True
                            
                            # Add the image
                            try:
                                doc.add_picture(img_path, width=Inches(4.0))
                            except Exception as img_e:
                                print(f"Warning: Could not add image {img_path} to Word doc: {img_e}")
                                p = doc.add_paragraph()
                                p.add_run(f"[Error loading image {image_index + 1}]")
                        else:
                            # Image not found
                            p = doc.add_paragraph()
                            p.add_run(f"[Image {image_index + 1} not found]")
                        
                        image_index += 1
                    except Exception as e:
                        print(f"Warning: Error processing image {image_index + 1}: {e}")
                        p = doc.add_paragraph()
                        p.add_run(f"[Error processing image {image_index + 1}]")
                        image_index += 1
                    continue
                
                # Add regular text line
                if line.strip() != "":
                    p = doc.add_paragraph(line)
                else:
                    # Add empty paragraph for empty lines
                    doc.add_paragraph()
            
            # Add footer with image count
            if image_index > 0:
                doc.add_paragraph()
                footer = doc.add_paragraph(f"Document exported with {image_index} image(s) | {time.strftime('%Y-%m-%d %H:%M:%S')}")
                footer.style.font.size = Pt(9)
                footer.style.font.color.rgb = RGBColor(128, 128, 128)
            
            doc.save(file_path)
            messagebox.showinfo("Export Successful", 
                              f"Document with images exported to:\n{file_path}\n\n"
                              f"Total images included: {image_index}")
            self.log_action(f"Exported Word document with images: {file_path}")
            
        except Exception as e:
            messagebox.showerror("Export Word", f"Failed to export: {e}")

    def export_excel(self):
        """Export text tab content to Excel file"""
        if not excel_available:
            messagebox.showerror("Export to Excel", 
                               "Excel export requires the 'openpyxl' library.\n\n"
                               "Install it using:\npip install openpyxl")
            return
            
        if not self.tabs:
            return
            
        tab = self.tabs[self.current_tab_index]
        
        if tab.is_drawing:
            messagebox.showinfo("Export to Excel", "Excel export is only available for text tabs")
            return
        
        file_path = filedialog.asksaveasfilename(
            defaultextension=".xlsx",
            filetypes=[("Excel Files", "*.xlsx")]
        )
        
        if not file_path:
            return
        
        try:
            # Create workbook and worksheet
            wb = Workbook()
            ws = wb.active
            # Title should be valid for Excel sheet names (max 31 chars, no certain symbols)
            ws.title = tab.name[:31].replace('[','').replace(']','').replace('*','').replace('?','').replace('/','').replace('\\','').replace(':','').replace('~','')
            if not ws.title: # If title becomes empty after cleaning
                ws.title = "Document"

            # Add header
            ws['A1'] = "Document Content"
            ws['A1'].font = ExcelFont(bold=True, size=14)
            ws['A1'].alignment = Alignment(horizontal='center')
            ws.merge_cells('A1:E1') # Merge cells for a wider header
            
            # Get text content
            content = tab.content.get("1.0", tk.END)
            
            # Add content starting from row 3 (after header)
            current_row = 3
            for line in content.split('\n'):
                # Check if the line is not just whitespace
                if line.strip():  
                    ws[f'A{current_row}'] = line
                    # Set alignment for wrap text and top vertical alignment
                    ws[f'A{current_row}'].alignment = Alignment(wrap_text=True, vertical='top')
                    current_row += 1
            
            # Adjust column width for better readability
            ws.column_dimensions['A'].width = 80 
            
            # Add metadata at the bottom
            current_row += 2 # Add some space between content and metadata
            ws[f'A{current_row}'] = "Document Statistics"
            ws[f'A{current_row}'].font = ExcelFont(bold=True, size=12)
            current_row += 1
            
            word_count = len(content.split())
            char_count = len(content)
            line_count = len(content.split('\n'))
            
            ws[f'A{current_row}'] = f"Word Count: {word_count}"
            current_row += 1
            ws[f'A{current_row}'] = f"Character Count: {char_count}"
            current_row += 1
            ws[f'A{current_row}'] = f"Line Count: {line_count}"
            current_row += 1
            ws[f'A{current_row}'] = f"Exported: {time.strftime('%Y-%m-%d %H:%M:%S')}"
            
            # Save the workbook
            wb.save(file_path)
            
            messagebox.showinfo("Export to Excel", f"Successfully exported to {file_path}")
            self.log_action(f"Exported to Excel: {file_path}")
            
        except Exception as e:
            messagebox.showerror("Export to Excel", f"Failed to export: {e}")

    def export_drawing_pdf(self):
        """Exports the current drawing tab content to a PDF file."""
        if not self.tabs: 
            return
        
        tab = self.tabs[self.current_tab_index]
        if not tab.is_drawing:
            messagebox.showinfo("Export Drawing PDF", "This function is only for drawing tabs.")
            return
        
        file_path = filedialog.asksaveasfilename(defaultextension=".pdf", 
                                               filetypes=[("PDF Files","*.pdf")])
        if not file_path:
            return
        
        try:
            # Use PIL to save the drawing to a temporary image file
            temp_img_path = os.path.join(tempfile.gettempdir(), 
                                        f"temp_drawing_export_{int(time.time())}.png")
            tab.content.save_image(temp_img_path)
            
            c = canvas.Canvas(file_path, pagesize=letter)
            
            # Add a title
            c.setFont("Helvetica-Bold", 16)
            c.drawString(100, 750, f"Drawing: {tab.name}")
            
            # Draw the image on the PDF
            pdf_width = letter[0] - 200  # Page width minus margins
            pdf_height = 500  # Fixed height for the image
            
            try:
                img_pil = Image.open(temp_img_path)
                img_w, img_h = img_pil.size
                
                # Calculate aspect ratio
                aspect_ratio = img_w / img_h
                draw_w = pdf_width
                draw_h = draw_w / aspect_ratio
                
                if draw_h > pdf_height:
                    draw_h = pdf_height
                    draw_w = draw_h * aspect_ratio
                
                # Center the image on the page
                x_offset = (letter[0] - draw_w) / 2
                y_offset = 600 - draw_h  # Position below title
                
                c.drawImage(temp_img_path, x_offset, y_offset, 
                           width=draw_w, height=draw_h)
                
                # Add a caption
                c.setFont("Helvetica", 10)
                c.drawString(100, y_offset - 20, f"Drawing saved: {time.strftime('%Y-%m-%d %H:%M:%S')}")
                
            except Exception as img_e:
                print(f"Warning: Could not draw image to PDF: {img_e}")
                c.drawString(100, 600, "[Drawing image could not be loaded]")
            
            # Add footer with metadata
            c.setFont("Helvetica", 8)
            c.drawString(100, 50, f"Exported from Notepad with Drawing | {time.strftime('%Y-%m-%d %H:%M:%S')}")
            
            c.save()
            messagebox.showinfo("Export Successful", 
                              f"Drawing exported to:\n{file_path}")
            self.log_action(f"Exported drawing to PDF: {file_path}")
            
        except Exception as e:
            messagebox.showerror("Export Drawing PDF", f"Failed to export: {e}")
        finally:
            # Clean up temporary image file
            if os.path.exists(temp_img_path):
                try:
                    os.remove(temp_img_path)
                except Exception as rm_e:
                    print(f"Warning: Could not remove temporary file {temp_img_path}: {rm_e}")

    def export_drawing_word(self):
        """Exports the current drawing tab content to a Word (.docx) file."""
        if not self.tabs: 
            return
        
        tab = self.tabs[self.current_tab_index]
        if not tab.is_drawing:
            messagebox.showinfo("Export Drawing Word", "This function is only for drawing tabs.")
            return
        
        file_path = filedialog.asksaveasfilename(defaultextension=".docx", 
                                               filetypes=[("Word Files","*.docx")])
        if not file_path:
            return
        
        try:
            doc = Document()
            
            # Add title
            title = doc.add_heading(f"Drawing: {tab.name}", 0)
            title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            
            # Save the current drawing to a temporary image file
            temp_img_path = os.path.join(tempfile.gettempdir(), 
                                        f"temp_drawing_export_{int(time.time())}.png")
            tab.content.save_image(temp_img_path)
            
            # Add the drawing image to the document
            try:
                # Add image with caption
                doc.add_paragraph("Drawing:", style='Heading 2')
                doc.add_picture(temp_img_path, width=Inches(6.0))
                
                # Add image caption
                caption = doc.add_paragraph(f"Drawing created on: {time.strftime('%Y-%m-%d %H:%M:%S')}")
                caption.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                caption.style.font.size = Pt(9)
                caption.style.font.italic = True
                
            except Exception as img_e:
                print(f"Warning: Could not add drawing image to Word doc: {img_e}")
                doc.add_paragraph("[Drawing image could not be loaded]", style='Body Text')
            
            # Add footer with metadata
            doc.add_paragraph()  # Add spacing
            footer = doc.add_paragraph(f"Exported from Notepad with Drawing | {time.strftime('%Y-%m-%d %H:%M:%S')}")
            footer.style.font.size = Pt(8)
            footer.style.font.color.rgb = RGBColor(128, 128, 128)  # Gray color
            
            doc.save(file_path)
            messagebox.showinfo("Export Successful", 
                              f"Drawing exported to:\n{file_path}")
            self.log_action(f"Exported drawing to Word: {file_path}")
            
        except Exception as e:
            messagebox.showerror("Export Drawing Word", f"Failed to export: {e}")
        finally:
            # Clean up temporary image file
            if os.path.exists(temp_img_path):
                try:
                    os.remove(temp_img_path)
                except Exception as rm_e:
                    print(f"Warning: Could not remove temporary file {temp_img_path}: {rm_e}")

    def export_image_with_current_text(self):
        """Exports an image with the current text tab content (from Edit menu)."""
        if not self.tabs or self.current_tab_index >= len(self.tabs):
            return
        
        current_tab = self.tabs[self.current_tab_index]
        if current_tab.is_drawing:
            messagebox.showinfo("Export Image with Text", "This function is for text tabs. Please select a text tab.")
            return
        
        # Find drawing tabs to choose from
        drawing_tabs = [tab for tab in self.tabs if tab.is_drawing]
        if not drawing_tabs:
            messagebox.showinfo("No Drawing Tabs", "No drawing tabs available. Please create a drawing tab first.")
            return
        
        # Create a dialog to select which drawing tab to use
        select_window = tk.Toplevel(self.root)
        select_window.title("Select Drawing")
        select_window.geometry("400x300")
        select_window.transient(self.root)
        select_window.grab_set()
        
        tk.Label(select_window, text="Select a drawing to combine with your text:", 
                font=("Arial", 11)).pack(pady=10)
        
        listbox_frame = tk.Frame(select_window)
        listbox_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=5)
        
        scrollbar = tk.Scrollbar(listbox_frame)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        
        tab_listbox = tk.Listbox(listbox_frame, yscrollcommand=scrollbar.set, 
                                font=("Arial", 10), selectmode=tk.SINGLE)
        tab_listbox.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        scrollbar.config(command=tab_listbox.yview)
        
        for i, tab in enumerate(drawing_tabs):
            tab_listbox.insert(tk.END, f"{tab.name}")
        
        def export_with_selected_drawing():
            try:
                selected_index = tab_listbox.curselection()[0]
                selected_drawing_tab = drawing_tabs[selected_index]
                
                # Store the link between text and drawing tabs
                current_tab.linked_drawing_tab_name = selected_drawing_tab.name
                selected_drawing_tab.linked_text_tab_name = current_tab.name
                selected_drawing_tab.drawing_text = current_tab.content.get("1.0", "end-1c")
                
                # Ask which format to export
                format_window = tk.Toplevel(self.root)
                format_window.title("Export Format")
                format_window.geometry("300x150")
                format_window.transient(self.root)
                format_window.grab_set()
                
                tk.Label(format_window, text="Choose export format:", 
                        font=("Arial", 11)).pack(pady=10)
                
                tk.Button(format_window, text="PDF", width=20,
                         command=lambda: [select_window.destroy(), format_window.destroy(), 
                                         self.export_drawing_pdf_with_selected(selected_drawing_tab)]).pack(pady=5)
                
                tk.Button(format_window, text="Word Document", width=20,
                         command=lambda: [select_window.destroy(), format_window.destroy(), 
                                         self.export_drawing_word_with_selected(selected_drawing_tab)]).pack(pady=5)
                
                tk.Button(format_window, text="Cancel", width=20,
                         command=lambda: [select_window.destroy(), format_window.destroy()]).pack(pady=5)
                
            except IndexError:
                messagebox.showwarning("No Selection", "Please select a drawing tab.")
        
        button_frame = tk.Frame(select_window)
        button_frame.pack(pady=10)
        
        tk.Button(button_frame, text="Export with Selected", command=export_with_selected_drawing).pack(side=tk.LEFT, padx=5)
        tk.Button(button_frame, text="Cancel", command=select_window.destroy).pack(side=tk.LEFT, padx=5)

    def export_drawing_pdf_with_selected(self, drawing_tab):
        """Exports the selected drawing tab with current text tab content to PDF."""
        file_path = filedialog.asksaveasfilename(defaultextension=".pdf", 
                                               filetypes=[("PDF Files","*.pdf")])
        if not file_path:
            return
        
        try:
            # Use PIL to save the drawing to a temporary image file
            temp_img_path = os.path.join(tempfile.gettempdir(), 
                                        f"temp_drawing_export_{int(time.time())}.png")
            drawing_tab.content.save_image(temp_img_path)
            
            c = canvas.Canvas(file_path, pagesize=letter)
            
            # Get current text tab
            current_tab = self.tabs[self.current_tab_index]
            
            # Add a title
            c.setFont("Helvetica-Bold", 16)
            c.drawString(100, 750, f"Document with Drawing: {drawing_tab.name}")
            
            c.setFont("Helvetica", 12)
            c.drawString(100, 730, f"Text from: {current_tab.name}")
            
            # Draw the image on the PDF
            pdf_width = letter[0] - 200  # Page width minus margins
            pdf_height = 400  # Fixed height for the image
            
            try:
                img_pil = Image.open(temp_img_path)
                img_w, img_h = img_pil.size
                
                # Calculate aspect ratio
                aspect_ratio = img_w / img_h
                draw_w = pdf_width
                draw_h = draw_w / aspect_ratio
                
                if draw_h > pdf_height:
                    draw_h = pdf_height
                    draw_w = draw_h * aspect_ratio
                
                # Center the image on the page
                x_offset = (letter[0] - draw_w) / 2
                y_offset = 600 - draw_h  # Position below title
                
                c.drawImage(temp_img_path, x_offset, y_offset, 
                           width=draw_w, height=draw_h)
                
                # Add a caption
                c.setFont("Helvetica", 10)
                c.drawString(100, y_offset - 20, f"Drawing saved: {time.strftime('%Y-%m-%d %H:%M:%S')}")
                
            except Exception as img_e:
                print(f"Warning: Could not draw image to PDF: {img_e}")
                c.drawString(100, 600, "[Drawing image could not be loaded]")
            
            # Add text content
            if hasattr(drawing_tab, 'drawing_text') and drawing_tab.drawing_text:
                # Add text header
                c.setFont("Helvetica-Bold", 14)
                c.drawString(100, 550, f"Text from {current_tab.name}:")
                
                # Prepare text
                text_content = drawing_tab.drawing_text
                c.setFont("Helvetica", 11)
                
                # Set up text parameters
                text_object = c.beginText(100, 520)
                text_object.setTextOrigin(100, 520)
                text_object.setFont("Helvetica", 11)
                
                # Split text into lines that fit the page width
                max_width = letter[0] - 200  # 100px margins on both sides
                line_height = 14
                
                for paragraph in text_content.split('\n'):
                    if paragraph.strip() == '':
                        text_object.textLine('')  # Empty line for paragraph break
                        continue
                    
                    # Handle long lines by word wrapping
                    words = paragraph.split()
                    current_line = ''
                    
                    for word in words:
                        test_line = current_line + (' ' if current_line else '') + word
                        # Simple width estimation (approximate)
                        if len(test_line) * 6 <= max_width:  # Rough estimate
                            current_line = test_line
                        else:
                            if current_line:
                                text_object.textLine(current_line)
                            current_line = word
                    
                    if current_line:
                        text_object.textLine(current_line)
                
                c.drawText(text_object)
            
            # Add footer with metadata
            c.setFont("Helvetica", 8)
            c.drawString(100, 50, f"Exported from Notepad with Drawing | {time.strftime('%Y-%m-%d %H:%M:%S')}")
            
            c.save()
            messagebox.showinfo("Export Successful", 
                              f"Drawing with text exported to:\n{file_path}")
            self.log_action(f"Exported drawing '{drawing_tab.name}' with text '{current_tab.name}' to PDF: {file_path}")
            
        except Exception as e:
            messagebox.showerror("Export PDF", f"Failed to export: {e}")
        finally:
            # Clean up temporary image file
            if os.path.exists(temp_img_path):
                try:
                    os.remove(temp_img_path)
                except Exception as rm_e:
                    print(f"Warning: Could not remove temporary file {temp_img_path}: {rm_e}")

    def export_drawing_word_with_selected(self, drawing_tab):
        """Exports the selected drawing tab with current text tab content to Word."""
        file_path = filedialog.asksaveasfilename(defaultextension=".docx", 
                                               filetypes=[("Word Files","*.docx")])
        if not file_path:
            return
        
        try:
            doc = Document()
            
            # Get current text tab
            current_tab = self.tabs[self.current_tab_index]
            
            # Add title
            title = doc.add_heading(f"Document with Drawing: {drawing_tab.name}", 0)
            title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            
            subtitle = doc.add_paragraph(f"Text from: {current_tab.name}")
            subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            subtitle.style.font.italic = True
            
            # Save the drawing to a temporary image file
            temp_img_path = os.path.join(tempfile.gettempdir(), 
                                        f"temp_drawing_export_{int(time.time())}.png")
            drawing_tab.content.save_image(temp_img_path)
            
            # Add the drawing image to the document
            try:
                # Add image with caption
                doc.add_paragraph("Drawing:", style='Heading 2')
                doc.add_picture(temp_img_path, width=Inches(6.0))
                
                # Add image caption
                caption = doc.add_paragraph(f"Drawing created on: {time.strftime('%Y-%m-%d %H:%M:%S')}")
                caption.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                caption.style.font.size = Pt(9)
                caption.style.font.italic = True
                
            except Exception as img_e:
                print(f"Warning: Could not add drawing image to Word doc: {img_e}")
                doc.add_paragraph("[Drawing image could not be loaded]", style='Body Text')
            
            # Add text content
            if hasattr(drawing_tab, 'drawing_text') and drawing_tab.drawing_text:
                doc.add_paragraph()  # Add spacing
                doc.add_heading(f"Text from {current_tab.name}", level=2)
                
                # Add the text content
                text_para = doc.add_paragraph(drawing_tab.drawing_text)
                text_para.style.font.size = Pt(11)
            
            # Add footer with metadata
            doc.add_paragraph()  # Add spacing
            footer = doc.add_paragraph(f"Exported from Notepad with Drawing | {time.strftime('%Y-%m-%d %H:%M:%S')}")
            footer.style.font.size = Pt(8)
            footer.style.font.color.rgb = RGBColor(128, 128, 128)  # Gray color
            
            doc.save(file_path)
            messagebox.showinfo("Export Successful", 
                              f"Drawing with text exported to:\n{file_path}")
            self.log_action(f"Exported drawing '{drawing_tab.name}' with text '{current_tab.name}' to Word: {file_path}")
            
        except Exception as e:
            messagebox.showerror("Export Word", f"Failed to export: {e}")
        finally:
            # Clean up temporary image file
            if os.path.exists(temp_img_path):
                try:
                    os.remove(temp_img_path)
                except Exception as rm_e:
                    print(f"Warning: Could not remove temporary file {temp_img_path}: {rm_e}")

    def enable_tab_reordering(self):
        """Enables drag-and-drop reordering of tabs."""
        self.notebook.bind("<ButtonPress-1>", self.on_tab_press)
        self.notebook.bind("<B1-Motion>", self.on_tab_motion)

    def on_tab_press(self, event):
        """Records the starting index when a tab is pressed for reordering."""
        try:
            self._drag_start_index = self.notebook.index(f"@{event.x},{event.y}")
        except tk.TclError:
            self._drag_start_index = None # Clicked outside of a tab

    def on_tab_motion(self, event):
        """Handles the motion of dragging a tab for reordering."""
        if self._drag_start_index is None:
            return
        try:
            # Get the index where the mouse is currently hovering over a tab
            new_index = self.notebook.index(f"@{event.x},{event.y}")
        except tk.TclError:
            return # Mouse is not over a tab

        if new_index != self._drag_start_index:
            tab_id = self.notebook.tabs()[self._drag_start_index] # Get the internal tab ID
            
            # Move the tab in the notebook widget
            self.notebook.insert(new_index, tab_id)
            
            # Update the order in our internal tabs list to match the notebook
            self.tabs.insert(new_index, self.tabs.pop(self._drag_start_index))
            
            # Update the drag start index to the new position
            self._drag_start_index = new_index

    def pick_theme(self):
        """Opens a dialog to select a theme for the current text tab."""
        theme_window = tk.Toplevel(self.root)
        theme_window.title("Choose Theme")
        theme_window.geometry("300x250")
        theme_window.transient(self.root) # Keep on top of main window
        theme_window.grab_set() # Modal dialog

        tk.Label(theme_window, text="Select a theme for this tab:", font=("Arial", 11)).pack(pady=10)
        
        themes = [
            ("Light", "light"),
            ("Dark", "dark"),
            ("Sepia", "sepia"),
            ("Blue", "blue"),
            ("Green", "green")
        ]
        
        def apply_theme(theme):
            """Applies the selected theme to the current text tab."""
            if self.current_tab_index < len(self.tabs):
                tab = self.tabs[self.current_tab_index]
                tab.theme = theme # Store the theme in the tab object
                
                if tab.is_drawing:
                    tab.content.set_theme(theme) # Apply theme to drawing canvas
                else:
                    # Define background and foreground colors for each theme
                    theme_bg = {
                        "light": "#ffffff",
                        "dark": "#1e1e1e",
                        "sepia": "#f4eae0",
                        "blue": "#e3f2fd",
                        "green": "#e8f5e9"
                    }
                    theme_fg = {
                        "light": "#000000",
                        "dark": "#ffffff",
                        "sepia": "#3e2723",
                        "blue": "#0d47a1",
                        "green": "#1b5e20"
                    }
                    # Apply colors to the text widget
                    tab.content.config(bg=theme_bg.get(theme), fg=theme_fg.get(theme), insertbackground=theme_fg.get(theme))
                self.log_action(f"Applied theme '{theme}' to tab '{tab.name}'")
                theme_window.destroy() # Close the theme selection window
        
        # Create buttons for each theme
        for theme_name, theme_val in themes:
            tk.Button(theme_window, text=theme_name, command=lambda t=theme_val: apply_theme(t), width=20).pack(pady=5)

    def search_text(self):
        """Opens a find dialog to search for text within the current tab."""
        if not self.tabs or self.current_tab_index >= len(self.tabs):
            return
        tab = self.tabs[self.current_tab_index]
        if tab.is_drawing:
            messagebox.showinfo("Search", "Cannot search in drawing tabs")
            return
        
        # Create the find dialog window
        dialog = tk.Toplevel(self.root)
        dialog.title("Find")
        dialog.geometry("300x150")
        dialog.transient(self.root)
        dialog.grab_set()
        
        tk.Label(dialog, text="Find text:").pack(pady=5)
        search_entry = tk.Entry(dialog, width=30)
        search_entry.pack(pady=5)
        search_entry.focus() # Set focus to the entry field
        
        def find_text():
            """Performs the search and highlights matches."""
            search_term = search_entry.get()
            if not search_term:
                return
            
            # Remove previous search highlights
            try:
                tab.content.tag_remove("search", "1.0", tk.END)
            except tk.TclError:
                pass # Tag might not exist
            
            idx = "1.0" # Start search from the beginning of the text
            found_count = 0
            while True:
                # Search for the term, case-insensitive, stopping at the end of the text
                idx = tab.content.search(search_term, idx, nocase=1, stopindex=tk.END)
                if not idx: # If no more matches are found
                    break
                found_count += 1
                end_idx = f"{idx}+{len(search_term)}c" # Calculate end index of the found text
                tab.content.tag_add("search", idx, end_idx) # Add tag for highlighting
                idx = end_idx # Move search start position past the current match
            
            tab.content.tag_config("search", background="#ffff00", foreground="black") # Configure highlight style
            
            if found_count > 0:
                self.log_action(f"Found {found_count} instance(s) of '{search_term}'")
                result_label.config(text=f"Found: {found_count} instance(s)", fg="green")
            else:
                result_label.config(text="Not found", fg="red")
        
        def clear_search():
            """Removes all search highlights."""
            try:
                tab.content.tag_remove("search", "1.0", tk.END)
                result_label.config(text="Search cleared", fg="blue")
                self.log_action("Cleared search highlights")
            except tk.TclError:
                pass # Tag might not exist
        
        button_frame = tk.Frame(dialog)
        button_frame.pack(pady=5)
        tk.Button(button_frame, text="Find All", command=find_text).pack(side=tk.LEFT, padx=5)
        tk.Button(button_frame, text="Clear", command=clear_search).pack(side=tk.LEFT, padx=5)
        
        result_label = tk.Label(dialog, text="", fg="black")
        result_label.pack(pady=5)
        
        # Bind Enter key to find_text function for convenience
        search_entry.bind("<Return>", lambda e: find_text())

    def undo(self):
        """Performs undo operation on the current text tab."""
        if not self.tabs: return
        tab = self.tabs[self.current_tab_index]
        if not tab.is_drawing:
            try:
                tab.content.edit_undo()
                self.log_action("Undo")
            except tk.TclError: # Ignore if no undo history
                pass

    def redo(self):
        """Performs redo operation on the current text tab."""
        if not self.tabs: return
        tab = self.tabs[self.current_tab_index]
        if not tab.is_drawing:
            try:
                tab.content.edit_redo()
                self.log_action("Redo")
            except tk.TclError: # Ignore if no redo history
                pass

    def cut(self):
        """Cuts selected text to the clipboard."""
        if not self.tabs: return
        tab = self.tabs[self.current_tab_index]
        if not tab.is_drawing:
            try:
                text = tab.content.get("sel.first", "sel.last")
                self.root.clipboard_clear()
                self.root.clipboard_append(text)
                tab.content.delete("sel.first", "sel.last")
                self.log_action("Cut")
            except tk.TclError: # No text selected
                pass

    def copy(self):
        """Copies selected text to the clipboard."""
        if not self.tabs: return
        tab = self.tabs[self.current_tab_index]
        if not tab.is_drawing:
            try:
                text = tab.content.get("sel.first", "sel.last")
                self.root.clipboard_clear()
                self.root.clipboard_append(text)
                self.log_action("Copy")
            except tk.TclError: # No text selected
                pass

    def paste(self):
        """Pastes text from the clipboard into the current text tab."""
        if self.tabs and self.current_tab_index < len(self.tabs):
            tab = self.tabs[self.current_tab_index]
            if not tab.is_drawing:
                # Tkinter's Text widget has built-in paste functionality via event generation
                tab.content.event_generate("<<Paste>>")
                self.log_action("Pasted text from clipboard")

    def delete_selection(self):
        """Deletes the selected text."""
        if self.tabs and self.current_tab_index < len(self.tabs):
            tab = self.tabs[self.current_tab_index]
            if not tab.is_drawing:
                try:
                    tab.content.delete(tk.SEL_FIRST, tk.SEL_LAST)
                    self.log_action("Deleted selected text")
                except tk.TclError: # No text selected
                    pass

    def print_file(self):
        """Simulates printing, currently prompts user to use Export PDF."""
        if self.tabs and self.current_tab_index < len(self.tabs):
            tab = self.tabs[self.current_tab_index]
            if not tab.is_drawing:
                content = tab.content.get("1.0", tk.END)
                messagebox.showinfo("Print", "Print functionality would normally send this document to your printer.\n\nFor now, please use 'Export PDF' to save and print from a PDF viewer.")
                self.log_action("Print requested for text tab")
            else:
                messagebox.showinfo("Print", "For drawing tabs, please use 'Export Drawing PDF' to save and print.")
                self.log_action("Print requested for drawing tab")

    def select_all(self):
        """Selects all text in the current text tab."""
        if self.tabs and self.current_tab_index < len(self.tabs):
            tab = self.tabs[self.current_tab_index]
            if not tab.is_drawing:
                tab.content.tag_add(tk.SEL, "1.0", tk.END) # Select all text
                tab.content.mark_set(tk.INSERT, "1.0")   # Move cursor to beginning
                tab.content.see(tk.INSERT)              # Scroll to cursor
                self.log_action("Selected all text")

    def show_highlight_palette(self):
        """Shows a palette of colors to highlight selected text."""
        if not self.tabs or self.current_tab_index >= len(self.tabs):
            return
        tab = self.tabs[self.current_tab_index]
        if not tab.is_drawing:
            try:
                # Check if any text is selected
                tab.content.index(tk.SEL_FIRST)
                
                palette_window = tk.Toplevel(self.root)
                palette_window.title("Choose Highlight Color")
                palette_window.geometry("350x280")
                palette_window.transient(self.root) # Modal
                palette_window.grab_set()
                
                tk.Label(palette_window, text="Select highlight color:", font=("Arial", 10, "bold")).pack(pady=10)
                
                # Predefined colors with names
                colors = [
                    ("Yellow", "#FFFF00"),
                    ("Green", "#00FF00"),
                    ("Cyan", "#00FFFF"),
                    ("Pink", "#FF69B4"),
                    ("Orange", "#FFA500"),
                    ("Light Blue", "#ADD8E6")
                ]
                
                button_frame = tk.Frame(palette_window)
                button_frame.pack(pady=10)
                
                # Create buttons for predefined colors
                for i, (name, color) in enumerate(colors):
                    row = i // 2
                    col = i % 2
                    btn = tk.Button(button_frame, text=name, bg=color, width=14, height=2,
                                  command=lambda c=color: self.apply_highlight(c, palette_window))
                    btn.grid(row=row, column=col, padx=5, pady=5)
                
                # Frame for custom color and remove highlight buttons
                bottom_frame = tk.Frame(palette_window)
                bottom_frame.pack(pady=5)
                
                custom_btn = tk.Button(bottom_frame, text="Custom Color", width=18, height=2,
                                      command=lambda: self.pick_custom_highlight_color(palette_window))
                custom_btn.pack(side=tk.LEFT, padx=5)
                
                transparent_btn = tk.Button(bottom_frame, text="Transparent\n(Remove Highlight)", width=18, height=2,
                                          command=lambda: self.remove_highlight(palette_window))
                transparent_btn.pack(side=tk.LEFT, padx=5)
                
            except tk.TclError: # No text selected
                messagebox.showinfo("No Selection", "Please select text to highlight")

    def pick_custom_highlight_color(self, window):
        """Opens color chooser for custom highlight color."""
        color = colorchooser.askcolor(title="Choose Highlight Color")[1]
        if color:
            self.apply_highlight(color, window) # Apply the chosen color

    def apply_highlight(self, color, window=None):
        """Applies the specified highlight color to the selected text."""
        if not self.tabs: return
        tab = self.tabs[self.current_tab_index]
        if not tab.is_drawing:
            try:
                sel_start = tab.content.index(tk.SEL_FIRST)
                sel_end = tab.content.index(tk.SEL_LAST)
                
                # Generate a unique tag name to avoid conflicts
                tag_name = f"highlight_{sel_start}_{sel_end}_{int(time.time()*1000)}" 
                tab.content.tag_add(tag_name, sel_start, sel_end)
                tab.content.tag_config(tag_name, background=color)
                
                if window: # Close the palette window if it was open
                    window.destroy()
                self.log_action(f"Applied highlight color {color}")
            except tk.TclError: # No text selected
                pass

    def remove_highlight(self, window=None):
        """Removes highlighting from the selected text."""
        if not self.tabs: return
        tab = self.tabs[self.current_tab_index]
        if not tab.is_drawing:
            try:
                sel_start = tab.content.index(tk.SEL_FIRST)
                sel_end = tab.content.index(tk.SEL_LAST)
                
                # Iterate through all tags applied to the selection
                for tag in tab.content.tag_names(sel_start):
                    # Check if it's a highlight tag we manage
                    if tag.startswith("highlight_"):
                        tag_ranges = tab.content.tag_ranges(tag)
                        if tag_ranges:
                            # Iterate through all ranges associated with this tag (a tag can span multiple non-contiguous areas)
                            for i in range(0, len(tag_ranges), 2):
                                tag_start = str(tag_ranges[i])
                                tag_end = str(tag_ranges[i+1])
                                # Check if the current selection overlaps with this tag's range
                                if (tab.content.compare(sel_start, "<=", tag_end) and 
                                    tab.content.compare(sel_end, ">=", tag_start)):
                                    # Remove the tag from the overlapping portion of the selection
                                    tab.content.tag_remove(tag, sel_start, sel_end)
                
                if window: # Close the palette window if it was open
                    window.destroy()
                self.log_action("Removed highlight")
            except tk.TclError: # No text selected
                pass

    def highlight_syntax(self, text_widget):
        """Basic syntax highlighting (currently only for spell check errors)."""
        if spell_checker is None:
            return
        
        try:
            text_widget.tag_remove("spelling", "1.0", tk.END) # Remove previous spell error tags
        except:
            pass # Tag might not exist
        
        content = text_widget.get("1.0", tk.END)
        
        try:
            # Check for spelling errors using enchant.check (returns list of tuples: (error_index, error_length))
            # NOTE: enchant.check is not directly compatible with Tkinter's get() in this way.
            # A more robust way is to iterate through words or use a different approach.
            # The current spell_checker.check(word) approach is more reliable.
            # We will use the real-time check method for individual words.
            pass # This function is now more of a placeholder for potential future syntax highlighting
        except Exception as e:
            pass # Ignore errors during potential syntax highlighting

    def change_font(self):
        """Applies the selected font family and size to the current text tab."""
        if not self.tabs: return
        tab = self.tabs[self.current_tab_index]
        if not tab.is_drawing:
            self.apply_font_style(tab) # Re-apply font style with new settings
            self.log_action(f"Changed font to {self.font_var.get()}")

    def change_font_size(self):
        """Applies the selected font size to the current text tab."""
        if not self.tabs: return
        tab = self.tabs[self.current_tab_index]
        if not tab.is_drawing:
            self.apply_font_style(tab) # Re-apply font style with new settings
            self.log_action(f"Changed font size to {self.size_var.get()}")

    def apply_font_style(self, tab):
        """Applies the current font, size, bold, italic, and underline settings to a tab."""
        if tab.is_drawing:
            return
        font_name = self.font_var.get()
        font_size = int(self.size_var.get())
        
        # Fetch current style settings from the toolbar buttons
        font_weight = "bold" if self.bold else "normal"
        font_slant = "italic" if self.italic else "roman"
        font_underline = 1 if self.underline else 0
        
        try:
            # Create a Tkinter Font object
            font = Font(family=font_name, size=font_size, weight=font_weight,
                        slant=font_slant, underline=font_underline)
            # Configure the text widget's font
            tab.content.config(font=font)
        except Exception as e:
            print(f"Error applying font style: {e}") # Log any font errors

    def toggle_bold(self):
        """Toggles bold formatting for the selected text."""
        if not self.tabs: return
        tab = self.tabs[self.current_tab_index]
        if not tab.is_drawing:
            try:
                # Get current selection
                sel_start = tab.content.index(tk.SEL_FIRST)
                sel_end = tab.content.index(tk.SEL_LAST)
                
                # Get current font properties from the selection or default
                current_font = self.get_selection_font(tab.content)
                
                # Toggle bold weight
                new_weight = "normal" if current_font.actual()["weight"] == "bold" else "bold"
                
                # Create new font with toggled weight
                new_font = Font(
                    family=current_font.actual()["family"],
                    size=current_font.actual()["size"],
                    weight=new_weight,
                    slant=current_font.actual()["slant"],
                    underline=current_font.actual()["underline"]
                )
                
                # Apply new font format to the selection
                self.apply_format_to_selection(tab.content, new_font, sel_start, sel_end)
                self.log_action("Toggled bold")
            except tk.TclError: # No text selected
                pass

    def toggle_italic(self):
        """Toggles italic formatting for the selected text."""
        if not self.tabs: return
        tab = self.tabs[self.current_tab_index]
        if not tab.is_drawing:
            try:
                sel_start = tab.content.index(tk.SEL_FIRST)
                sel_end = tab.content.index(tk.SEL_LAST)
                
                current_font = self.get_selection_font(tab.content)
                
                # Toggle italic slant
                new_slant = "roman" if current_font.actual()["slant"] == "italic" else "italic"
                
                new_font = Font(
                    family=current_font.actual()["family"],
                    size=current_font.actual()["size"],
                    weight=current_font.actual()["weight"],
                    slant=new_slant,
                    underline=current_font.actual()["underline"]
                )
                
                self.apply_format_to_selection(tab.content, new_font, sel_start, sel_end)
                self.log_action("Toggled italic")
            except tk.TclError: # No text selected
                pass

    def toggle_underline(self):
        """Toggles underline formatting for the selected text."""
        if not self.tabs: return
        tab = self.tabs[self.current_tab_index]
        if not tab.is_drawing:
            try:
                sel_start = tab.content.index(tk.SEL_FIRST)
                sel_end = tab.content.index(tk.SEL_LAST)
                
                current_font = self.get_selection_font(tab.content)
                
                # Toggle underline property
                new_underline = 0 if current_font.actual()["underline"] else 1
                
                new_font = Font(
                    family=current_font.actual()["family"],
                    size=current_font.actual()["size"],
                    weight=current_font.actual()["weight"],
                    slant=current_font.actual()["slant"],
                    underline=new_underline
                )
                
                self.apply_format_to_selection(tab.content, new_font, sel_start, sel_end)
                self.log_action("Toggled underline")
            except tk.TclError: # No text selected
                pass

    def get_selection_font(self, text_widget):
        """Retrieves the font properties of the current selection, or defaults to the widget's font."""
        base_font = Font(font=text_widget.cget("font"))
        try:
            # Check if there's an active formatting tag on the selection
            # We'll check the first character for existing formatting tags
            sel_start = text_widget.index(tk.SEL_FIRST)
            current_tags = text_widget.tag_names(sel_start)
            
            for tag in current_tags:
                if tag.startswith("format_"): # Our custom formatting tag
                    font_config = text_widget.tag_cget(tag, "font")
                    if font_config:
                        return Font(font=font_config) # Return specific font config if found
            
            # If no specific formatting tag found, return the base font
            return base_font
        except tk.TclError: # No text selected or other error
            return base_font # Return default font

    def apply_format_to_selection(self, text_widget, font, sel_start, sel_end):
        """Applies a specific font format to the selected text using a unique tag."""
        # Generate a unique tag name to avoid conflicts and ensure distinct formatting
        tag_name = f"format_{sel_start}_{sel_end}_{int(time.time()*1000)}" 
        
        # Remove any previous format tags on the selection to avoid overwriting issues
        # This is a simplified approach; a more robust solution might involve merging font properties
        current_tags = text_widget.tag_names(sel_start)
        for tag in current_tags:
            if tag.startswith("format_"):
                text_widget.tag_remove(tag, sel_start, sel_end)

        # Add the new formatting tag
        text_widget.tag_add(tag_name, sel_start, sel_end)
        text_widget.tag_config(tag_name, font=font)

    def update_style_buttons(self):
        """Updates the visual state of the Bold, Italic, Underline buttons based on selection."""
        if not self.tabs: return
        tab = self.tabs[self.current_tab_index]
        if tab.is_drawing:
            # Disable style buttons for drawing tabs
            self.bold_btn.config(state=tk.DISABLED, bg="#cccccc")
            self.italic_btn.config(state=tk.DISABLED, bg="#cccccc")
            self.underline_btn.config(state=tk.DISABLED, bg="#cccccc")
            self.highlight_btn.config(state=tk.DISABLED, bg="#cccccc")
            if self.spell_check_enabled:
                self.spell_check_btn.config(state=tk.DISABLED, bg="#cccccc")
            return
        else:
            # Enable style buttons for text tabs
            self.bold_btn.config(state=tk.NORMAL)
            self.italic_btn.config(state=tk.NORMAL)
            self.underline_btn.config(state=tk.NORMAL)
            self.highlight_btn.config(state=tk.NORMAL)
            if self.spell_check_enabled:
                self.spell_check_btn.config(state=tk.NORMAL)

        try:
            # Get font properties of the current selection
            current_font = self.get_selection_font(tab.content)
            
            # Update self.bold, self.italic, self.underline based on current font properties
            self.bold = current_font.actual()["weight"] == "bold"
            self.italic = current_font.actual()["slant"] == "italic"
            self.underline = current_font.actual()["underline"] == 1
        except tk.TclError: # No selection, use default state or global settings
            # This part might need refinement if global styles are different from default
            pass # Keep previous state if no selection

        # Update button background colors
        self.bold_btn.config(bg="#ffeb3b" if self.bold else "#f3f3f3")
        self.italic_btn.config(bg="#ffeb3b" if self.italic else "#f3f3f3")
        self.underline_btn.config(bg="#ffeb3b" if self.underline else "#f3f3f3")

    def autosave(self):
        """Saves the current state of all open tabs periodically."""
        for i, tab in enumerate(self.tabs):
            if tab.is_drawing:
                # Auto-save drawing to its last saved path if available
                if hasattr(tab, 'auto_save_path') and tab.auto_save_path:
                    try:
                        tab.content.save_image(tab.auto_save_path)
                        self.log_action(f"Auto-saved drawing: {tab.name}")
                    except:
                        pass # Ignore errors during autosave
            else:
                # Auto-save text file to its original path if it has one
                if tab.file_path:
                    try:
                        with open(tab.file_path, "w", encoding="utf-8") as f:
                            f.write(tab.content.get("1.0", tk.END))
                        self.log_action(f"Auto-saved: {tab.file_path}")
                    except:
                        pass # Ignore errors during autosave
        # Schedule the next autosave
        self.root.after(30000, self.autosave)

    def toggle_dark_mode(self):
        """Toggles dark mode for the entire application."""
        self.dark_mode = not self.dark_mode
        for tab in self.tabs:
            if tab.is_drawing:
                tab.content.set_dark_mode(self.dark_mode) # Apply dark mode to drawing canvas
            else:
                # Apply dark mode colors to text widget
                bg = "#1e1e1e" if self.dark_mode else "#ffffff"
                fg = "#ffffff" if self.dark_mode else "#000000"
                try:
                    tab.content.config(bg=bg, fg=fg, insertbackground=fg)
                except:
                    pass # Ignore errors if tab content is not yet initialized

        # Update toolbar background to match dark mode
        toolbar_bg = "#252526" if self.dark_mode else "#f3f3f3"
        self.root.children['!frame'].config(bg=toolbar_bg) # Assuming toolbar is the first frame
        for widget in self.root.children['!frame'].winfo_children():
            if isinstance(widget, tk.Frame): # Style the main toolbar frame
                widget.config(bg=toolbar_bg)
                for sub_widget in widget.winfo_children():
                    if isinstance(sub_widget, tk.Label):
                        sub_widget.config(bg=toolbar_bg, fg="#ffffff" if self.dark_mode else "#000000")
                    elif isinstance(sub_widget, (tk.Button, ttk.Combobox)):
                        sub_widget.config(bg=toolbar_bg, fg="#ffffff" if self.dark_mode else "#000000")
                    elif isinstance(sub_widget, tk.Scale):
                        sub_widget.config(bg=toolbar_bg, troughcolor="#555555" if self.dark_mode else "#cccccc")
                    elif isinstance(sub_widget, tk.Radiobutton):
                        sub_widget.config(bg=toolbar_bg, fg="#ffffff" if self.dark_mode else "#000000",
                                          activebackground="#666666" if self.dark_mode else "#cccccc",
                                          selectcolor="#777777" if self.dark_mode else "#a0a0a0")
            elif isinstance(widget, tk.Label): # Style the status label
                widget.config(bg=toolbar_bg, fg="#000000" if self.dark_mode else "#000000") # Status bar text color might need adjustment
        
        self.log_action(f"Dark mode {'enabled' if self.dark_mode else 'disabled'}")

    def log_action(self, action):
        """Adds an action to the activity log."""
        timestamp = time.strftime("%Y-%m-%d %H:%M:%S")
        self.activity_log.append(f"[{timestamp}] {action}")
        # Keep log size manageable
        if len(self.activity_log) > 100:
            self.activity_log.pop(0)

    def open_activity_log(self):
        """Opens a new window displaying the activity log."""
        log_window = tk.Toplevel(self.root)
        log_window.title("Activity Log")
        log_window.geometry("600x400")
        log_window.transient(self.root)

        text_widget = tk.Text(log_window, wrap=tk.WORD, state=tk.NORMAL, font=("Arial", 9))
        text_widget.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)
        
        # Insert log entries
        for log_entry in self.activity_log:
            text_widget.insert(tk.END, log_entry + "\n")
        
        text_widget.config(state=tk.DISABLED) # Make read-only

    def manage_personal_dictionary(self):
        """Opens a window to manage the personal dictionary."""
        dict_window = tk.Toplevel(self.root)
        dict_window.title("Personal Dictionary")
        dict_window.geometry("400x500")
        dict_window.transient(self.root)
        dict_window.grab_set() # Modal

        tk.Label(dict_window, text="Words in your personal dictionary:", font=("Arial", 11, "bold")).pack(pady=10)

        list_frame = tk.Frame(dict_window)
        list_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=5)

        scrollbar = tk.Scrollbar(list_frame)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)

        self.dictionary_listbox = tk.Listbox(list_frame, yscrollcommand=scrollbar.set, font=("Arial", 10), selectmode=tk.SINGLE)
        self.dictionary_listbox.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        scrollbar.config(command=self.dictionary_listbox.yview)

        # Populate the listbox with words from the personal_dictionary set
        for word in sorted(personal_dictionary):
            self.dictionary_listbox.insert(tk.END, word)

        # Input field for adding new words
        add_frame = tk.Frame(dict_window)
        add_frame.pack(pady=5)
        
        new_word_entry = tk.Entry(add_frame, width=30, font=("Arial", 10))
        new_word_entry.pack(side=tk.LEFT, padx=5)
        new_word_entry.focus()

        def add_word_to_dict():
            """Adds a word from the entry field to the personal dictionary."""
            word = new_word_entry.get().strip().lower()
            if word and word.isalpha(): # Ensure it's a valid word (only alphabetic chars)
                if word not in personal_dictionary:
                    personal_dictionary.add(word)
                    # Add to listbox and save
                    self.dictionary_listbox.insert(tk.END, word)
                    save_personal_dictionary()
                    self.log_action(f"Added '{word}' to personal dictionary.")
                    messagebox.showinfo("Success", f"'{word}' added.")
                    new_word_entry.delete(0, tk.END) # Clear the entry field
                else:
                    messagebox.showinfo("Info", f"'{word}' is already in the dictionary.")
            else:
                messagebox.showwarning("Invalid Input", "Please enter a valid word (alphabetic characters only).")

        add_btn = tk.Button(add_frame, text="Add Word", command=add_word_to_dict)
        add_btn.pack(side=tk.LEFT)
        new_word_entry.bind("<Return>", lambda e: add_word_to_dict()) # Add on Enter key press

        # Button to remove selected word
        remove_btn = tk.Button(dict_window, text="Remove Selected Word", command=self.remove_selected_word)
        remove_btn.pack(pady=5)
        
        # Close button
        close_btn = tk.Button(dict_window, text="Close", command=dict_window.destroy)
        close_btn.pack(pady=10)

    def remove_selected_word(self):
        """Removes the word selected in the listbox from the personal dictionary."""
        try:
            selected_index = self.dictionary_listbox.curselection()[0] # Get index of selected item
            word_to_remove = self.dictionary_listbox.get(selected_index) # Get word from index
            
            if word_to_remove in personal_dictionary:
                personal_dictionary.remove(word_to_remove)
                self.dictionary_listbox.delete(selected_index) # Remove from listbox
                save_personal_dictionary() # Save updated dictionary
                self.log_action(f"Removed '{word_to_remove}' from personal dictionary.")
                messagebox.showinfo("Success", f"'{word_to_remove}' removed.")
        except IndexError: # No item selected
            messagebox.showwarning("No Selection", "Please select a word to remove.")

    def save_session(self):
        """Saves the current state of tabs and their contents to a JSON file."""
        session_data = {
            "used_text_tab_numbers": list(self.used_text_tab_numbers),
            "used_drawing_tab_numbers": list(self.used_drawing_tab_numbers),
            "tabs": [] # List to store info about each tab
        }
        for i, tab in enumerate(self.tabs):
            tab_info = {
                "tab_number": getattr(tab, 'tab_number', None) # Store the unique tab number
            }
            if tab.is_drawing:
                # Save drawing to a temporary file
                img_path = os.path.join(tempfile.gettempdir(), f"session_drawing_{i}_{int(time.time())}.png")
                try:
                    tab.content.save_image(img_path)
                    tab_info.update({
                        "type": "drawing",
                        "path": img_path, # Path to the saved drawing image
                        "name": tab.name,
                        "drawing_text": tab.drawing_text if hasattr(tab, 'drawing_text') else "", # Save associated text if any
                        "linked_text_tab_name": tab.linked_text_tab_name if hasattr(tab, 'linked_text_tab_name') else ""
                    })
                    session_data["tabs"].append(tab_info)
                except Exception as e:
                    print(f"Error saving drawing tab to session: {e}")
            else:
                # Save text content to a temporary file
                text_path = os.path.join(tempfile.gettempdir(), f"session_text_{i}_{int(time.time())}.txt")
                try:
                    with open(text_path, "w", encoding="utf-8") as f:
                        f.write(tab.content.get("1.0", tk.END))
                    tab_info.update({
                        "type": "text",
                        "path": text_path, # Path to the saved text file
                        "name": tab.name,
                        "file_path": tab.file_path # Original file path if it was opened/saved
                    })
                    session_data["tabs"].append(tab_info)
                except Exception as e:
                    print(f"Error saving text tab to session: {e}")
        
        # Save session data to JSON file
        try:
            with open(self.session_file, "w") as f:
                json.dump(session_data, f)
            self.log_action("Session saved with all tabs")
        except Exception as e:
            print(f"Error saving session file: {e}")

    def load_session(self):
        """Loads the application state from the session file."""
        if not os.path.exists(self.session_file):
            return # No session file to load
        
        try:
            with open(self.session_file, "r") as f:
                session_data = json.load(f)
        except Exception as e:
            print(f"Error loading session file: {e}")
            return
        
        # Restore tab numbers
        self.used_text_tab_numbers = set(session_data.get("used_text_tab_numbers", []))
        self.used_drawing_tab_numbers = set(session_data.get("used_drawing_tab_numbers", []))
        
        # Remove the initial "Untitled" tab if it's still there and we're loading a session
        if len(self.tabs) == 1 and self.tabs[0].name.startswith("Untitled"):
            try:
                # Remove its tab number if it was assigned
                if hasattr(self.tabs[0], 'tab_number'):
                    self.used_text_tab_numbers.discard(self.tabs[0].tab_number)
                self.notebook.forget(0) # Remove from notebook widget
                self.tabs.pop(0)       # Remove from list
            except tk.TclError:
                pass # Ignore if it's already gone
        
        # Recreate tabs from session data
        for tab_info in session_data.get("tabs", []):
            try:
                if tab_info.get("type") == "text":
                    text_path = tab_info.get("path")
                    if os.path.exists(text_path):
                        with open(text_path, "r", encoding="utf-8") as f:
                            content = f.read()
                        
                        # Create a new Tab object
                        tab = Tab(tab_info.get("name", "Untitled"), False)
                        tab.file_path = tab_info.get("file_path") # Original file path
                        if tab_info.get("tab_number") is not None:
                            tab.tab_number = tab_info["tab_number"]
                        self.tabs.append(tab)
                        
                        # Create the Text widget and its frame
                        frame = tk.Frame(self.notebook)
                        text_widget = tk.Text(frame, wrap=tk.WORD, undo=True)
                        text_widget.pack(fill=tk.BOTH, expand=True)
                        text_widget.bind("<KeyRelease>", lambda e: (self.update_status_bar(e), self.highlight_syntax(text_widget), self.check_spelling_realtime(text_widget)))
                        text_widget.insert("1.0", content) # Load content
                        tab.content = text_widget
                        
                        self.notebook.add(frame, text=tab.name) # Add tab to notebook
                        self.log_action(f"Restored text tab: {tab.name}")
                
                elif tab_info.get("type") == "drawing":
                    img_path = tab_info.get("path")
                    if os.path.exists(img_path):
                        # Create a new Tab object
                        tab = Tab(tab_info.get("name", "Drawing"), True)
                        tab.drawing_text = tab_info.get("drawing_text", "") # Restore associated text
                        tab.linked_text_tab_name = tab_info.get("linked_text_tab_name", "") # Restore linked tab name
                        if tab_info.get("tab_number") is not None:
                            tab.tab_number = tab_info["tab_number"]
                        self.tabs.append(tab)
                        
                        # Create the DrawingCanvas and its frame
                        frame = tk.Frame(self.notebook)
                        drawing_canvas = DrawingCanvas(frame, self, self.dark_mode) # Pass current dark_mode
                        drawing_canvas.frame.pack(fill=tk.BOTH, expand=True)
                        
                        # Load the drawing from the temporary image file
                        try:
                            img = Image.open(img_path)
                            # Paste the loaded image onto the canvas's internal image object
                            # Resize to canvas dimensions if necessary (though ideally saved at canvas size)
                            img_resized = img.resize((drawing_canvas.canvas_width, drawing_canvas.canvas_height))
                            drawing_canvas.image.paste(img_resized)
                            # Redraw the canvas content (optional, if the image object itself isn't directly drawn)
                            # For this implementation, the DrawingCanvas's image attribute is the source of truth,
                            # so we might need a method to refresh the canvas display if it's not automatic.
                            # For now, we assume pasting into `drawing_canvas.image` is sufficient for future saves.
                            # If visual update on load is needed, a `redraw_canvas()` method would be required.
                        except Exception as load_img_e:
                            print(f"Error loading drawing image {img_path}: {load_img_e}")
                        
                        tab.content = drawing_canvas
                        self.notebook.add(frame, text=tab.name)
                        self.log_action(f"Restored drawing tab: {tab.name}")
            except Exception as e:
                print(f"Error restoring tab '{tab_info.get('name', 'Unknown')}': {e}")
        
        self.log_action("Session loaded successfully")

    def on_close(self):
        """Handles the window closing event, saving session and personal dictionary."""
        save_personal_dictionary() # Save user-added words
        self.save_session()       # Save current tabs and content
        
        # Ensure drawings are saved if they have an auto_save_path set
        for tab in self.tabs:
            if tab.is_drawing and hasattr(tab, 'auto_save_path') and tab.auto_save_path:
                try:
                    tab.content.save_image(tab.auto_save_path)
                    self.log_action(f"Saved drawing on exit: {tab.name}")
                except:
                    pass # Ignore errors during save on exit
        
        self.log_action("Application closed - session saved")
        self.root.destroy() # Close the main window

    def show_about(self):
        """Displays the 'About' message box."""
        messagebox.showinfo("About", "Notepad with Drawing\nVersion 2.0\n\nFeatures:\n- Text and Drawing tabs\n- Spell checking with PyEnchant\n- Persistent sessions\n- Theme support\n- Bold, Italic, Underline text formatting\n- Combined formatting support\n- Custom highlight colors\n- Smart tab numbering\n- Insert images into text documents\n- Export documents with images to PDF/Word\n- Personal Dictionary management\n- Export to Excel")

    def spell_check(self):
        """Performs a manual spell check on the selected text in a text tab."""
        if spell_checker:
            if not self.tabs: return
            tab = self.tabs[self.current_tab_index]
            if tab.is_drawing:
                messagebox.showinfo("Spell Check", "Cannot spell check in drawing tabs")
                return

            # Check if text is selected
            try:
                sel_start = tab.content.index(tk.SEL_FIRST)
                sel_end = tab.content.index(tk.SEL_LAST)
                content = tab.content.get(sel_start, sel_end)
            except tk.TclError:
                # If no text is selected, prompt user to select text
                messagebox.showinfo("Spell Check", "Please select text to spell check")
                return

            try:
                import re
                words = re.findall(r'\b[a-zA-Z]+\b', content) # Extract words from selection
                misspelled = []

                for word in words:
                    # Check if the word is not recognized by enchant and not in personal dictionary
                    if not spell_checker.check(word) and word.lower() not in personal_dictionary:
                        suggestions = spell_checker.suggest(word) # Get suggestions
                        suggestion_text = f" (suggestions: {', '.join(suggestions[:3])})" if suggestions else "" # Limit suggestions shown
                        misspelled.append(f"'{word}'{suggestion_text}")

                # Remove duplicate misspelled words while preserving order
                seen = set()
                unique_misspelled = []
                for item in misspelled:
                    if item not in seen:
                        seen.add(item)
                        unique_misspelled.append(item)

                if unique_misspelled:
                    # Display up to 10 misspelled words in a message box
                    issues_str = "\n".join(unique_misspelled[:10])
                    count_msg = f" (showing 10 of {len(unique_misspelled)})" if len(unique_misspelled) > 10 else ""
                    messagebox.showinfo("Spell Check", f"Found misspelled words in selection{count_msg}:\n\n{issues_str}")
                else:
                    messagebox.showinfo("Spell Check", "No spelling issues found in selection!")
            except Exception as e:
                messagebox.showerror("Spell Check", f"Error: {e}")

            self.log_action("Spell checked selected text")

    def check_spelling_realtime(self, text_widget):
        """Checks spelling in real-time as the user types and underlines misspelled words."""
        if not self.spell_check_enabled or not spell_checker:
            return
        
        try:
            # Remove all existing spell error tags before re-tagging
            text_widget.tag_remove("spell_error", "1.0", tk.END)
            
            # Get all text content from the widget
            content = text_widget.get("1.0", tk.END)
            
            # Use regex to find all words
            import re
            for match in re.finditer(r'\b[a-zA-Z]+\b', content):
                word = match.group()
                # Check if the word is misspelled and not in the personal dictionary
                if not spell_checker.check(word) and word.lower() not in personal_dictionary:
                    # Calculate the start and end positions for tagging
                    start_pos = f"1.0 + {match.start()} chars"
                    end_pos = f"1.0 + {match.end()} chars"
                    text_widget.tag_add("spell_error", start_pos, end_pos) # Apply the spell_error tag
        except Exception as e:
            pass  # Silently ignore errors during real-time spell checking

    def show_spell_context_menu(self, event, text_widget):
        """Shows a context menu with spell suggestions and an option to add to dictionary when right-clicking on a misspelled word."""
        if not self.spell_check_enabled or not spell_checker:
            return
        
        # Get the character index at the mouse cursor position
        index = text_widget.index(f"@{event.x},{event.y}")
        
        # Determine the boundaries of the word at that index
        word_start = text_widget.search(r'\s', index, backwards=True, regexp=True)
        if not word_start:
            word_start = "1.0" # If no whitespace before, it's the start of the text
        else:
            word_start = f"{word_start}+1c" # Move past the whitespace
        
        word_end = text_widget.search(r'\s', index, regexp=True)
        if not word_end:
            word_end = tk.END # If no whitespace after, it's the end of the text
        
        # Get the word itself and clean it (remove surrounding punctuation)
        word = text_widget.get(word_start, word_end).strip()
        
        import string
        word_clean = word.strip(string.punctuation)
        
        # If the cleaned word is not a valid alphabetic word, do nothing
        if not word_clean or not word_clean.isalpha():
            return
        
        # Check if the word is actually misspelled (not recognized and not in personal dictionary)
        if spell_checker.check(word_clean) or word_clean.lower() in personal_dictionary:
            return  # Word is spelled correctly or ignored, so no context menu needed
        
        # Create the context menu
        context_menu = tk.Menu(text_widget, tearoff=0)
        
        # Get spelling suggestions
        suggestions = spell_checker.suggest(word_clean)
        
        if suggestions:
            # Add suggestions to the context menu
            for i, suggestion in enumerate(suggestions[:10]):  # Limit to top 10 suggestions
                context_menu.add_command(
                    label=suggestion,
                    # Command to replace the misspelled word with the suggestion
                    command=lambda s=suggestion, ws=word_start, we=word_end: self.replace_word(text_widget, ws, we, s)
                )
            context_menu.add_separator()
        else:
            context_menu.add_command(label="(No suggestions)", state="disabled")
            context_menu.add_separator()
        
        # Add option to add the word to the personal dictionary
        context_menu.add_command(
            label=f"Add '{word_clean}' to Dictionary",
            command=lambda w=word_clean: self.add_to_personal_dictionary(w, text_widget)
        )
        
        # Add option to ignore the word (add to personal dictionary for this session)
        context_menu.add_command(
            label=f"Ignore All",
            command=lambda w=word_clean: self.ignore_word(w, text_widget)
        )
        
        # Display the context menu
        try:
            context_menu.tk_popup(event.x_root, event.y_root)
        finally:
            context_menu.grab_release()

    def replace_word(self, text_widget, start, end, new_word):
        """Replaces a word in the text widget with a new word."""
        old_word = text_widget.get(start, end)
        
        # Preserve any trailing punctuation or whitespace that was part of the original selection
        import string
        trailing = ""
        for char in reversed(old_word):
            if char in string.punctuation or char.isspace():
                trailing = char + trailing
            else:
                break
        
        # Delete the old word and insert the new one with preserved trailing characters
        text_widget.delete(start, end)
        text_widget.insert(start, new_word + trailing)
        
        self.check_spelling_realtime(text_widget) # Re-check spelling after modification
        self.log_action(f"Replaced word with suggestion: {new_word}")

    def add_to_personal_dictionary(self, word, text_widget):
        """Adds a word to the personal dictionary, saves it, and updates spell check."""
        personal_dictionary.add(word.lower()) # Add lowercase version
        save_personal_dictionary() # Save to file
        self.check_spelling_realtime(text_widget) # Update real-time highlighting
        self.log_action(f"Added '{word}' to personal dictionary")
        messagebox.showinfo("Dictionary", f"'{word}' added to personal dictionary")

    def ignore_word(self, word, text_widget):
        """Temporarily ignores a word by adding it to the personal dictionary for the current session."""
        personal_dictionary.add(word.lower()) # Add to the set (persists if session is saved)
        self.check_spelling_realtime(text_widget) # Update real-time highlighting
        self.log_action(f"Ignoring '{word}' in current session")

    def update_status_bar(self, event=None):
        """Updates the status bar with word count, character count, line, and column."""
        if not self.tabs: return
        tab = self.tabs[self.current_tab_index]
        if not tab.is_drawing:
            # Get text content excluding trailing newline character added by Text widget
            content = tab.content.get("1.0", "end-1c")
            word_count = len(content.split())
            char_count = len(content)
            
            # Get cursor position to display line and column number
            cursor_pos = tab.content.index("insert") # e.g., "5.23" for line 5, column 23
            line, column = map(int, cursor_pos.split('.'))
            
            # Calculate total lines
            line_count = int(tab.content.index("end-1c").split(".")[0])

            status_text = f"Words: {word_count} | Characters: {char_count} | Lines: {line_count} | Line: {line} | Column: {column}"
            self.toolbar_status.config(text=status_text)
        else:
            # Update status bar for drawing mode
            self.toolbar_status.config(text="Drawing Mode")

if __name__ == "__main__":
    root = tk.Tk()
    app = NotepadApp(root)
    root.mainloop()